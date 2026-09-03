"""Canonical release and container validation tooling for Milestone 7.

Validates the full Docker Compose deployment, LibreOffice/PDF execution,
network isolation, read-only template mount, and transient artifact lifecycle.

This is an authoritative, local production-like smoke validator using
DOMAIN=localhost and testing against host ports published by Docker Compose.
"""

from __future__ import annotations

import http.client
import json
from pathlib import Path
import socket
import ssl
import subprocess
import sys
import time
import urllib.request
import yaml

WORKSPACE_ROOT = Path(__file__).resolve().parent.parent


def run_cmd(cmd: list[str], check: bool = True, cwd: Path = WORKSPACE_ROOT) -> subprocess.CompletedProcess:
    """Run a shell command and capture output."""
    res = subprocess.run(cmd, cwd=str(cwd), capture_output=True, text=True)
    if check and res.returncode != 0:
        print(f"COMMAND FAILED: {' '.join(cmd)}", file=sys.stderr)
        print(f"STDOUT: {res.stdout}", file=sys.stderr)
        print(f"STDERR: {res.stderr}", file=sys.stderr)
        raise RuntimeError(f"Command failed with exit code {res.returncode}: {' '.join(cmd)}")
    return res


def validate_release() -> bool:
    print("================================================================================")
    print("DOCUMENT BUILDER — MILESTONE 7 RELEASE VALIDATION")
    print("================================================================================\n")

    # A. BUILD / CONFIG
    print("[A] Validating Build & Compose Config...")
    run_cmd(["docker", "compose", "config"])
    print("  ✔ docker compose config validates cleanly")

    run_cmd(["docker", "compose", "build"])
    print("  ✔ docker compose build succeeded")

    run_cmd(["docker", "compose", "up", "-d"])
    print("  ✔ docker compose up -d executed")

    # B. HEALTH & REVERSE PROXY
    print("\n[B] Verifying App Health & Reverse Proxy (HTTP -> HTTPS)...")
    # Poll until app container Health.Status is EXACTLY "healthy" (no substring matching)
    healthy = False
    last_status = "unknown"
    for _ in range(20):
        # Resolve container ID for app service
        cid_res = run_cmd(["docker", "compose", "ps", "-q", "app"], check=False)
        cid = cid_res.stdout.strip()
        if cid:
            insp = run_cmd(["docker", "inspect", "--format", "{{.State.Health.Status}}", cid], check=False)
            last_status = insp.stdout.strip()
            if last_status == "healthy":
                healthy = True
                break
        time.sleep(1.5)

    if not healthy:
        raise RuntimeError(f"App container did not become healthy in time. Last status: {last_status!r}")
    print(f"  ✔ app container Health.Status is EXACTLY 'healthy' (status={last_status})")

    # Resolve published Caddy ports dynamically from docker compose
    def resolve_caddy_port(container_port: int) -> tuple[str, int]:
        p_res = run_cmd(["docker", "compose", "port", "caddy", str(container_port)])
        line = p_res.stdout.strip().splitlines()[0]
        if line.startswith("["):
            host_part, port_part = line.rsplit(":", 1)
            raw_host = host_part.strip("[]")
        else:
            raw_host, port_part = line.rsplit(":", 1)
        port_num = int(port_part)
        test_host = "127.0.0.1" if raw_host in ("0.0.0.0", "::", "") else raw_host
        return test_host, port_num

    http_host, http_port = resolve_caddy_port(80)
    https_host, https_port = resolve_caddy_port(443)
    print(f"  ✔ Resolved published Caddy ports: HTTP={http_host}:{http_port}, HTTPS={https_host}:{https_port}")

    # Explicit smoke check of HTTP endpoint without following redirects
    http_conn = http.client.HTTPConnection(http_host, http_port, timeout=5)
    try:
        http_conn.request("GET", "/health")
        http_resp = http_conn.getresponse()
        if http_resp.status not in (301, 307, 308):
            raise RuntimeError(f"Expected HTTP redirect (301/307/308), got {http_resp.status}")
        location = http_resp.getheader("Location")
        if not location or not location.startswith("https://"):
            raise RuntimeError(f"Expected HTTPS Location header from HTTP redirect, got {location}")
        print(f"  ✔ HTTP endpoint received status {http_resp.status} redirecting to {location}")
    finally:
        http_conn.close()

    # Request through Caddy over HTTPS
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE

    https_url = f"https://localhost:{https_port}/health" if https_port != 443 else "https://localhost/health"
    req = urllib.request.Request(https_url, headers={"User-Agent": "ReleaseValidation/1.0"})
    with urllib.request.urlopen(req, context=ctx, timeout=10) as resp:
        if resp.status != 200:
            raise RuntimeError(f"Expected 200 from Caddy HTTPS, got {resp.status}")
        data = json.loads(resp.read().decode("utf-8"))
        if data.get("status") != "ok":
            raise RuntimeError(f"Unexpected health response: {data}")
    print(f"  ✔ Caddy HTTPS {https_url} returned 200 and {{'status': 'ok'}}")

    # C. NETWORK EXPOSURE
    print("\n[C] Verifying Network Isolation & Host Ports...")
    # Port 8000 should NOT be directly reachable on host
    s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    s.settimeout(1.0)
    result = s.connect_ex(("127.0.0.1", 8000))
    s.close()
    if result == 0:
        raise RuntimeError("SECURITY VIOLATION: app port 8000 is directly exposed to host!")
    print("  ✔ app:8000 is NOT published to host (connection refused as expected)")

    # Caddy HTTPS published port should be open and reachable
    s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    s.settimeout(1.0)
    result = s.connect_ex((https_host, https_port))
    s.close()
    if result != 0:
        raise RuntimeError(f"Caddy HTTPS port {https_port} is not reachable on host {https_host}!")
    print(f"  ✔ Caddy host port {https_port} is published and accepting connections")


    # D. TEMPLATE MOUNT
    print("\n[D] Verifying Template Mount & Read-Only Enforcement...")
    test_tmpl_cmd = [
        "docker", "compose", "exec", "-T", "app", "python3", "-c",
        """
import os
from pathlib import Path
from hashoej_document_builder.core.discovery import discover_templates, discover_enabled_templates

root = Path("templates")
all_tmpls = discover_templates(root)
enabled = discover_enabled_templates(root)

assert len(all_tmpls) == 21, f"Expected 21 templates, found {len(all_tmpls)}"
assert len(enabled) == 20, f"Expected 20 enabled templates, found {len(enabled)}"

hif17 = next(t for t in all_tmpls if t.id == "hif-17-incident")
assert hif17.enabled is False, "HIF-17 must remain disabled"

# Test read-only enforcement
try:
    (root / "probe.txt").write_text("test")
    raise AssertionError("Write to templates succeeded! Mount is NOT read-only.")
except OSError:
    pass  # Expected read-only file system

print("OK")
"""
    ]
    res = run_cmd(test_tmpl_cmd)
    if "OK" not in res.stdout:
        raise RuntimeError(f"Template mount verification failed: {res.stdout} {res.stderr}")
    print("  ✔ All 21 TemplatePackages visible inside container")
    print("  ✔ Template directory mount is strictly read-only (write probe rejected)")
    print("  ✔ hif-17-incident is verified disabled (privacy/legal gate active)")
    print("  ✔ Public discovery returns exactly 20 enabled templates")

    # E. LIBREOFFICE & RUNTIME FONTS
    print("\n[E] Verifying LibreOffice Headless & Runtime Fonts...")
    test_lo_cmd = [
        "docker", "compose", "exec", "-T", "app", "python3", "-c",
        """
import shutil, tempfile, subprocess
from pathlib import Path

soffice = shutil.which("soffice") or shutil.which("libreoffice")
assert soffice is not None, "soffice binary not found in container"

with tempfile.TemporaryDirectory(prefix="lo_test_") as prof:
    uri = Path(prof).resolve().as_uri()
    res = subprocess.run([soffice, f"-env:UserInstallation={uri}", "--headless", "--version"], capture_output=True, text=True)
    assert res.returncode == 0, f"LibreOffice failed: {res.stderr}"

# Verify fontconfig resolution for Aptos and Aptos Display
fc_aptos = subprocess.run(["fc-match", "Aptos"], capture_output=True, text=True).stdout.strip()
fc_aptos_disp = subprocess.run(["fc-match", "Aptos Display"], capture_output=True, text=True).stdout.strip()
assert "Carlito" in fc_aptos, f"Aptos did not resolve to Carlito: {fc_aptos}"
assert "Carlito" in fc_aptos_disp, f"Aptos Display did not resolve to Carlito: {fc_aptos_disp}"

print(f"OK: {res.stdout.strip()} | Aptos->{fc_aptos}")
"""
    ]
    res = run_cmd(test_lo_cmd)
    if "OK" not in res.stdout:
        raise RuntimeError(f"LibreOffice check failed: {res.stdout} {res.stderr}")
    print("  ✔ soffice/libreoffice binary found in app image")
    print("  ✔ Isolated-profile headless invocation succeeds")
    print("  ✔ Aptos and Aptos Display fontconfig resolution verified to Carlito")

    # F & G. REPRESENTATIVE PRODUCTION-LIKE GENERATION & CLEANUP
    print("\n[F & G] Verifying Representative Generation & Transient Cleanup...")
    # Load profile data from host and pass to container
    normal_profile_text = (WORKSPACE_ROOT / "tests" / "compatibility" / "profiles" / "hif-01-role" / "normal.yaml").read_text(encoding="utf-8")

    test_gen_cmd = [
        "docker", "compose", "exec", "-T", "app", "python3", "-c",
        f"""
import docx, os, subprocess, time, yaml
from pathlib import Path
from hashoej_document_builder.core.discovery import load_template_package
from hashoej_document_builder.core.validation import validate_all_steps_values
from hashoej_document_builder.core.rendering import build_render_context, render_docx
from hashoej_document_builder.core.pdf import convert_docx_to_pdf
from hashoej_document_builder.core.artifacts import ArtifactManager

pkg = load_template_package(Path("templates/hif-01-role"))
raw_profile = yaml.safe_load({json.dumps(normal_profile_text)})

val_res = validate_all_steps_values(pkg.form_definition, raw_profile)
assert val_res.is_valid, f"Validation errors: {{val_res.errors}}"

art_mgr = ArtifactManager()
docx_path = art_mgr.create_artifact_path("docx")
render_ctx = build_render_context(pkg.form_definition, val_res.coerced_values)
render_docx(pkg.document_template, render_ctx, docx_path)

assert docx_path.is_file() and docx_path.stat().st_size > 0
doc = docx.Document(str(docx_path))
all_text = " ".join(p.text for p in doc.paragraphs) + " " + " ".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
assert "Ungdomsformand Fodbold" in all_text

pdf_path = convert_docx_to_pdf(docx_path, docx_path.parent)
assert pdf_path.is_file() and pdf_path.stat().st_size > 0
assert pdf_path.read_bytes()[:5] == b"%PDF-"

# Verify embedded fonts via pdffonts
fonts_res = subprocess.run(["pdffonts", str(pdf_path)], capture_output=True, text=True).stdout
assert "Carlito" in fonts_res, f"Expected Carlito embedded in PDF, got: {{fonts_res}}"
assert "DejaVuSerif" not in fonts_res, f"Unexpected DejaVuSerif found in PDF: {{fonts_res}}"

# Cleanup verification
art_mgr.cleanup_artifact(docx_path.parent)
assert not docx_path.parent.exists(), "Artifact folder was not removed"

# Stale cleanup proof:
# 1. Create dedicated transient artifact directory/file through existing ArtifactManager
stale_file = art_mgr.create_artifact_path("docx", prefix="stale_proof")
stale_dir = stale_file.parent
stale_file.write_bytes(b"stale_test_data")
assert stale_dir.is_dir()

# 2. Make it genuinely older than configured stale threshold (15 minutes in the past)
old_timestamp = time.time() - (15 * 60)
os.utime(str(stale_dir), (old_timestamp, old_timestamp))
os.utime(str(stale_file), (old_timestamp, old_timestamp))

# Create a fresh artifact to ensure unrelated paths are NOT touched
fresh_file = art_mgr.create_artifact_path("docx", prefix="fresh_proof")
fresh_dir = fresh_file.parent
fresh_file.write_bytes(b"fresh_test_data")
assert fresh_dir.is_dir()

# 3. Invoke cleanup_stale_artifacts with 10-minute threshold
purged_count = art_mgr.cleanup_stale_artifacts(max_age_minutes=10)

# 4. Assert stale artifact directory was actually removed
assert not stale_dir.exists(), "Stale artifact directory was not removed"
assert purged_count >= 1, f"Expected at least 1 purged directory, got {{purged_count}}"

# 5. Verify unrelated fresh path was untouched
assert fresh_dir.is_dir(), "Fresh artifact directory was incorrectly removed"
art_mgr.cleanup_artifact(fresh_dir)
assert not fresh_dir.exists()

print(f"OK: docx_size={{docx_path}}, pdf_size={{pdf_path.stat().st_size if pdf_path.exists() else 'cleaned'}}, purged={{purged_count}}")
"""
    ]
    res = run_cmd(test_gen_cmd)
    if "OK" not in res.stdout:
        raise RuntimeError(f"Representative generation failed: {res.stdout} {res.stderr}")
    print("  ✔ Representative TemplatePackage (hif-01-role) loaded and validated")
    print("  ✔ Authoritative DOCX rendered, reopened, and text verified")
    print("  ✔ Authoritative PDF converted and header validated (%PDF-)")
    print("  ✔ PDF font embedding verified (Carlito embedded, DejaVuSerif absent)")
    print("  ✔ Transient artifact directory explicitly cleaned up")
    print("  ✔ Stale-artifact purge verified with genuine aged directory and fresh isolation")

    print("\n================================================================================")
    print("ALL CANONICAL RELEASE GATES PASSED (A through G)")
    print("================================================================================\n")
    return True


if __name__ == "__main__":
    success = validate_release()
    if not success:
        sys.exit(1)
