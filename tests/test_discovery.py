from pathlib import Path
import docx
import pytest
import yaml

from hashoej_document_builder.core.discovery import (
    discover_enabled_templates,
    discover_templates,
    load_template_package,
)
from hashoej_document_builder.core.errors import (
    TemplateMissingFileError,
    TemplateNotFoundError,
    TemplateValidationError,
    TemplateYAMLError,
)


def _create_synthetic_template(
    dir_path: Path,
    template_id: str = "hif-01-role",
    version: int = 1,
    enabled: bool = True,
    title: str = "Rollebeskrivelse",
    category: str = "Organisation",
    description: str = "Beskriv en rolle.",
    include_docx: bool = True,
    yaml_content: str | None = None,
    docx_text: str = "{{ role_name }}",
) -> Path:
    dir_path.mkdir(parents=True, exist_ok=True)

    if yaml_content is None:
        raw_dict = {
            "id": template_id,
            "version": version,
            "enabled": enabled,
            "title": title,
            "category": category,
            "description": description,
            "steps": [
                {
                    "id": "basic",
                    "title": "Grundoplysninger",
                    "fields": [
                        {
                            "id": "role_name",
                            "type": "text",
                            "label": "Rollenavn",
                            "required": True,
                        }
                    ],
                }
            ],
        }
        yaml_content = yaml.dump(raw_dict)

    (dir_path / "template.yaml").write_text(yaml_content, encoding="utf-8")

    if include_docx:
        doc = docx.Document()
        doc.add_paragraph(docx_text)
        doc.save(str(dir_path / "document.docx"))

    return dir_path


def test_load_valid_template_package(tmp_path: Path) -> None:
    pkg_dir = tmp_path / "hif-01-role"
    _create_synthetic_template(pkg_dir)

    package = load_template_package(pkg_dir)

    assert package.id == "hif-01-role"
    assert package.version == 1
    assert package.enabled is True
    assert package.title == "Rollebeskrivelse"
    assert package.category == "Organisation"
    assert package.description == "Beskriv en rolle."
    assert package.path == pkg_dir.resolve()
    assert package.document_template == (pkg_dir / "document.docx").resolve()
    assert isinstance(package.form_definition, dict)
    assert package.form_definition["id"] == "hif-01-role"


def test_load_disabled_template_package(tmp_path: Path) -> None:
    pkg_dir = tmp_path / "hif-02-task"
    _create_synthetic_template(pkg_dir, template_id="hif-02-task", enabled=False, title="Opgavekort")

    package = load_template_package(pkg_dir)

    assert package.id == "hif-02-task"
    assert package.enabled is False


def test_discover_templates_returns_all_and_enabled_filters(tmp_path: Path) -> None:
    root = tmp_path / "templates"
    root.mkdir()

    # Create 3 templates out of order: hif-07 (enabled), hif-02 (disabled), hif-01 (enabled)
    _create_synthetic_template(root / "dir_c", template_id="hif-07-event", title="Arrangementsskabelon", enabled=True)
    _create_synthetic_template(root / "dir_b", template_id="hif-02-task", title="Opgavekort", enabled=False)
    _create_synthetic_template(root / "dir_a", template_id="hif-01-role", title="Rollebeskrivelse", enabled=True)

    all_packages = discover_templates(root)
    assert len(all_packages) == 3
    # Check deterministic ordering by id
    assert [p.id for p in all_packages] == ["hif-01-role", "hif-02-task", "hif-07-event"]

    enabled_packages = discover_enabled_templates(root)
    assert len(enabled_packages) == 2
    assert [p.id for p in enabled_packages] == ["hif-01-role", "hif-07-event"]


def test_discover_duplicate_template_id_same_version_rejected(tmp_path: Path) -> None:
    root = tmp_path / "templates"
    root.mkdir()

    _create_synthetic_template(root / "pkg_a", template_id="hif-01-role", version=1)
    _create_synthetic_template(root / "pkg_b", template_id="hif-01-role", version=1)

    with pytest.raises(TemplateValidationError, match="Duplicate template id 'hif-01-role'"):
        discover_templates(root)


def test_discover_duplicate_template_id_different_version_rejected(tmp_path: Path) -> None:
    root = tmp_path / "templates"
    root.mkdir()

    _create_synthetic_template(root / "pkg_v1", template_id="hif-01-role", version=1)
    _create_synthetic_template(root / "pkg_v2", template_id="hif-01-role", version=2)

    with pytest.raises(TemplateValidationError, match="Duplicate template id 'hif-01-role'"):
        discover_templates(root)


def test_discover_distinct_template_ids_remain_valid(tmp_path: Path) -> None:
    root = tmp_path / "templates"
    root.mkdir()

    _create_synthetic_template(root / "pkg_1", template_id="hif-01-role")
    _create_synthetic_template(root / "pkg_2", template_id="hif-02-task")
    _create_synthetic_template(root / "pkg_3", template_id="hif-03-policy")

    packages = discover_templates(root)
    assert len(packages) == 3
    assert [p.id for p in packages] == ["hif-01-role", "hif-02-task", "hif-03-policy"]


def test_discover_empty_directory(tmp_path: Path) -> None:
    root = tmp_path / "empty_templates"
    root.mkdir()

    packages = discover_templates(root)
    assert packages == []


def test_discover_non_existent_root_raises_not_found(tmp_path: Path) -> None:
    non_existent = tmp_path / "missing_dir"
    with pytest.raises(TemplateNotFoundError, match="Template root directory not found"):
        discover_templates(non_existent)


def test_load_non_existent_package_raises_not_found(tmp_path: Path) -> None:
    non_existent = tmp_path / "missing_pkg"
    with pytest.raises(TemplateNotFoundError, match="Template package directory not found"):
        load_template_package(non_existent)


def test_missing_template_yaml_raises_missing_file_error(tmp_path: Path) -> None:
    pkg_dir = tmp_path / "hif-01-role"
    pkg_dir.mkdir()
    doc = docx.Document()
    doc.save(str(pkg_dir / "document.docx"))

    with pytest.raises(TemplateMissingFileError, match="Missing required regular file 'template.yaml'"):
        load_template_package(pkg_dir)


def test_missing_document_docx_raises_missing_file_error(tmp_path: Path) -> None:
    pkg_dir = tmp_path / "hif-01-role"
    _create_synthetic_template(pkg_dir, include_docx=False)

    with pytest.raises(TemplateMissingFileError, match="Missing required regular file 'document.docx'"):
        load_template_package(pkg_dir)


def test_template_yaml_as_directory_rejected(tmp_path: Path) -> None:
    pkg_dir = tmp_path / "hif-01-role"
    pkg_dir.mkdir()
    # Create template.yaml as directory instead of regular file
    (pkg_dir / "template.yaml").mkdir()
    doc = docx.Document()
    doc.save(str(pkg_dir / "document.docx"))

    with pytest.raises(TemplateMissingFileError, match="Missing required regular file 'template.yaml'"):
        load_template_package(pkg_dir)


def test_document_docx_as_directory_rejected(tmp_path: Path) -> None:
    pkg_dir = tmp_path / "hif-01-role"
    pkg_dir.mkdir()
    (pkg_dir / "template.yaml").write_text("id: hif-01\n", encoding="utf-8")
    # Create document.docx as directory instead of regular file
    (pkg_dir / "document.docx").mkdir()

    with pytest.raises(TemplateMissingFileError, match="Missing required regular file 'document.docx'"):
        load_template_package(pkg_dir)


def test_invalid_yaml_syntax_raises_yaml_error(tmp_path: Path) -> None:
    pkg_dir = tmp_path / "hif-01-role"
    bad_yaml = "id: hif-01\nversion: [unclosed list"
    _create_synthetic_template(pkg_dir, yaml_content=bad_yaml)

    with pytest.raises(TemplateYAMLError, match="Invalid YAML"):
        load_template_package(pkg_dir)


def test_empty_yaml_raises_yaml_error(tmp_path: Path) -> None:
    pkg_dir = tmp_path / "hif-01-role"
    _create_synthetic_template(pkg_dir, yaml_content="")

    with pytest.raises(TemplateYAMLError, match="must be a non-empty YAML"):
        load_template_package(pkg_dir)


def test_invalid_schema_in_package_raises_validation_error(tmp_path: Path) -> None:
    pkg_dir = tmp_path / "hif-01-role"
    invalid_schema_yaml = """
id: hif-01-role
version: 1
enabled: true
title: Rollebeskrivelse
category: Organisation
description: Beskrivelse
steps:
  - id: basic
    title: Grundoplysninger
    fields:
      - id: role_name
        type: unknown_primitive_widget
        label: Rollenavn
"""
    _create_synthetic_template(pkg_dir, yaml_content=invalid_schema_yaml)

    with pytest.raises(TemplateValidationError, match="Unsupported field type 'unknown_primitive_widget'"):
        load_template_package(pkg_dir)
