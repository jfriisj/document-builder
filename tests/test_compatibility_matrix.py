"""Compatibility matrix and journey tests for HIF-01, HIF-02, and HIF-07 real templates."""

from __future__ import annotations

from pathlib import Path
import shutil
import docx
from fastapi.testclient import TestClient
import pytest
import yaml

from hashoej_document_builder.core.discovery import discover_templates, load_template_package
from hashoej_document_builder.core.docx_binding import validate_docx_binding
from hashoej_document_builder.core.pdf import convert_docx_to_pdf
from hashoej_document_builder.core.rendering import build_render_context, render_docx
from hashoej_document_builder.core.schema import validate_template_definition
from hashoej_document_builder.core.validation import validate_all_steps_values
from hashoej_document_builder.web.app import create_app

WORKSPACE_ROOT = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = WORKSPACE_ROOT / "templates"
PROFILES_DIR = WORKSPACE_ROOT / "tests" / "compatibility" / "profiles"

TEMPLATE_IDS = [
    "hif-01-role",
    "hif-02-task",
    "hif-07-event",
    "hif-03-handover",
    "hif-08-project",
    "hif-09-minutes",
    "hif-12-volunteer-shift",
    "hif-13-volunteer-onboarding",
    "hif-17-incident",
    "hif-20-communication",
    "hif-04-annual-cycle",
    "hif-06-contract",
    "hif-10-decision-log",
    "hif-11-action-list",
    "hif-15-inventory",
    "hif-16-maintenance",
    "hif-05-contact",
    "hif-14-key-access",
    "hif-18-purchase",
    "hif-19-sponsor",
    "hif-21-gdpr",
]
PROFILE_NAMES = ["minimal", "normal", "edge"]


def _extract_all_docx_text(doc: docx.Document) -> str:
    """Extract all text from paragraphs, tables, headers, and footers of a docx document."""
    text_parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            text_parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        text_parts.append(p.text)
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text:
                text_parts.append(p.text)
        for p in section.footer.paragraphs:
            if p.text:
                text_parts.append(p.text)
    return " ".join(text_parts)


@pytest.mark.parametrize("template_id", TEMPLATE_IDS)
def test_template_package_loads_and_is_enabled(template_id: str) -> None:
    """Verify that each real template package is valid, enabled, and discoverable."""
    pkg_dir = TEMPLATES_DIR / template_id
    assert pkg_dir.is_dir(), f"Template directory {pkg_dir} does not exist"

    pkg = load_template_package(pkg_dir)
    assert pkg.id == template_id
    if template_id == "hif-17-incident":
        # HIF-17 contains injury/medical data and must remain disabled until human privacy/legal gate
        assert pkg.enabled is False
    else:
        assert pkg.enabled is True
    assert len(pkg.form_definition["steps"]) >= 1

    # Discovery check
    all_pkgs = discover_templates(TEMPLATES_DIR)
    discovered_ids = [p.id for p in all_pkgs]
    assert template_id in discovered_ids


@pytest.mark.parametrize("template_id", TEMPLATE_IDS)
def test_template_contains_contextual_privacy_info(template_id: str) -> None:
    """Verify that templates processing personal/contact references provide at least one wizard privacy info field."""
    pkg_dir = TEMPLATES_DIR / template_id
    pkg = load_template_package(pkg_dir)
    form_def = pkg.form_definition

    privacy_info_fields = []
    for step in form_def.get("steps", []):
        for field in step.get("fields", []):
            if field.get("type") == "info" and field.get("variant") == "privacy":
                privacy_info_fields.append(field)

    assert len(privacy_info_fields) >= 1, (
        f"Template {template_id} must contain at least one field with type: info and variant: privacy"
    )
    for field in privacy_info_fields:
        text = field.get("text", "")
        assert len(text) > 0
        assert "session" in text.lower() or "permanent" in text.lower()


@pytest.mark.parametrize("template_id", TEMPLATE_IDS)
@pytest.mark.parametrize("profile_name", PROFILE_NAMES)
def test_compatibility_render_matrix(template_id: str, profile_name: str, tmp_path: Path) -> None:
    """Verify the 3x3 matrix: TemplatePackage loads, YAML validates, DOCX binding validates,

    values validate, render context builds, DOCX renders, opens, and contains expected values and structural rows.
    """
    pkg_dir = TEMPLATES_DIR / template_id
    profile_path = PROFILES_DIR / template_id / f"{profile_name}.yaml"
    assert profile_path.is_file(), f"Profile file {profile_path} does not exist"

    # 1. TemplatePackage loads & validates
    pkg = load_template_package(pkg_dir)
    form_def = pkg.form_definition

    # 2. Schema validation
    validated_def = validate_template_definition(form_def)
    assert validated_def["id"] == template_id

    # 3. DOCX binding validation
    validate_docx_binding(pkg.document_template, validated_def)

    # 4. Profile values validation
    raw_profile = yaml.safe_load(profile_path.read_text(encoding="utf-8"))
    assert isinstance(raw_profile, dict), f"Profile {profile_path} must be a dict"

    validation_result = validate_all_steps_values(validated_def, raw_profile)
    assert validation_result.is_valid, f"Validation errors for {template_id}/{profile_name}: {validation_result.errors}"

    # 5. Build render context
    context = build_render_context(validated_def, validation_result.coerced_values)
    assert isinstance(context, dict)

    # 6. Render DOCX
    out_docx = tmp_path / f"{template_id}_{profile_name}.docx"
    rendered_path = render_docx(pkg.document_template, context, out_docx)
    assert rendered_path.is_file()
    assert rendered_path.stat().st_size > 0

    # 7. Generated DOCX opens and contains tables
    doc = docx.Document(str(rendered_path))
    doc_text = _extract_all_docx_text(doc)

    # 8. Structural table row count assertions for normal / edge profiles
    if template_id == "hif-01-role":
        assert len(doc.tables) == 8  # Metadata + 7 numbered sections
        role_name = validation_result.coerced_values.get("role_name", "")
        if role_name:
            assert role_name in doc_text
        if profile_name in ("normal", "edge"):
            # Table 2: responsibilities (header + items)
            assert len(doc.tables[2].rows) >= 3
            # Table 4: access_resources (header + items)
            assert len(doc.tables[4].rows) >= 3
            # Table 5: annual_cycle (header + items)
            assert len(doc.tables[5].rows) >= 3

    elif template_id == "hif-02-task":
        assert len(doc.tables) == 10  # Metadata + 8 sections (section 7 has 2 tables)
        task_title = validation_result.coerced_values.get("task_title", "")
        if task_title:
            assert task_title in doc_text
        if profile_name in ("normal", "edge"):
            # Table 3: prerequisites (header + items)
            assert len(doc.tables[3].rows) >= 3
            # Table 4: materials (header + items)
            assert len(doc.tables[4].rows) >= 3
            # Table 5: task_steps (header + items)
            assert len(doc.tables[5].rows) >= 4
            # Table 7: completion_checks (header + items)
            assert len(doc.tables[7].rows) >= 3

    elif template_id == "hif-07-event":
        assert len(doc.tables) == 10  # Metadata + 9 numbered sections
        event_title = validation_result.coerced_values.get("event_title", "")
        if event_title:
            assert event_title in doc_text
        if profile_name in ("normal", "edge"):
            # Table 2: success_criteria
            assert len(doc.tables[2].rows) >= 3
            # Table 3: timeline
            assert len(doc.tables[3].rows) >= 4
            # Table 4: staffing
            assert len(doc.tables[4].rows) >= 4
            # Table 5: logistics
            assert len(doc.tables[5].rows) >= 4
            # Table 6: budget_items
            assert len(doc.tables[6].rows) >= 4
            # Table 7: communications
            assert len(doc.tables[7].rows) >= 3

    elif template_id == "hif-03-handover":
        assert len(doc.tables) == 9
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[2].rows) >= 3  # ongoing_tasks
            assert len(doc.tables[3].rows) >= 2  # recurring_tasks
            assert len(doc.tables[4].rows) >= 2  # access_and_systems

    elif template_id == "hif-08-project":
        assert len(doc.tables) == 9
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[3].rows) >= 2  # stakeholders
            assert len(doc.tables[4].rows) >= 2  # milestones
            assert len(doc.tables[5].rows) >= 2  # budget_items

    elif template_id == "hif-09-minutes":
        assert len(doc.tables) == 6
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[2].rows) >= 2  # agenda_items
            assert len(doc.tables[3].rows) >= 2  # action_items
            assert len(doc.tables[4].rows) >= 2  # logged_decisions

    elif template_id == "hif-12-volunteer-shift":
        assert len(doc.tables) == 7
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[3].rows) >= 2  # procedure_steps
            assert len(doc.tables[4].rows) >= 2  # equipment_access
            assert len(doc.tables[6].rows) >= 2  # closing_checks

    elif template_id == "hif-13-volunteer-onboarding":
        assert len(doc.tables) == 7
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[2].rows) >= 2  # club_info_topics
            assert len(doc.tables[4].rows) == 9  # fixed practical checks
            assert len(doc.tables[5].rows) >= 2  # onboarding_activities

    elif template_id == "hif-17-incident":
        assert len(doc.tables) == 8
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[3].rows) >= 2  # involved_persons
            assert len(doc.tables[5].rows) >= 2  # documentation_items
            assert len(doc.tables[6].rows) >= 2  # followup_actions

    elif template_id == "hif-20-communication":
        assert len(doc.tables) == 6
        assert len(doc.tables[4].rows) == 9  # 8 fixed quality checkpoints

    elif template_id == "hif-04-annual-cycle":
        assert len(doc.tables) == 4
        # Invariant: fixed 12-month scaffold across all profiles (minimal, normal, edge)
        assert len(doc.tables[1].rows) == 13
        expected_months = [
            "Januar", "Februar", "Marts", "April", "Maj", "Juni",
            "Juli", "August", "September", "Oktober", "November", "December",
        ]
        assert [r.cells[0].text.strip() for r in doc.tables[1].rows[1:]] == expected_months
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[2].rows) >= 3  # floating_deadlines (header + >=2 items)
            assert len(doc.tables[3].rows) == 4  # evaluation

    elif template_id == "hif-06-contract":
        assert len(doc.tables) == 4
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[1].rows) >= 3  # contracts (header + >=2 items)
            assert len(doc.tables[2].rows) >= 2  # renewal_checks
            assert len(doc.tables[3].rows) == 4  # general_terms
        if profile_name == "edge":
            assert len(doc.tables[2].rows) >= 3  # renewal_checks in edge has >= 2 items

    elif template_id == "hif-10-decision-log":
        assert len(doc.tables) == 2
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[1].rows) >= 3  # decisions (header + >=2 items)

    elif template_id == "hif-11-action-list":
        assert len(doc.tables) == 2
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[1].rows) >= 3  # action_items (header + >=2 items)

    elif template_id == "hif-15-inventory":
        assert len(doc.tables) == 3
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[1].rows) >= 3  # inventory_items (header + >=2 items)
            assert len(doc.tables[2].rows) >= 2  # loan_records
        if profile_name == "edge":
            assert len(doc.tables[2].rows) >= 3  # loan_records in edge has >= 2 items

    elif template_id == "hif-16-maintenance":
        assert len(doc.tables) == 4
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[1].rows) >= 3  # maintenance_tasks (header + >=2 items)
            assert len(doc.tables[2].rows) >= 2  # repairs_and_defects
            assert len(doc.tables[3].rows) == 3  # planning
        if profile_name == "edge":
            assert len(doc.tables[2].rows) >= 3  # repairs_and_defects in edge has >= 2 items

    elif template_id == "hif-05-contact":
        assert len(doc.tables) == 3
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[1].rows) >= 3  # contacts (header + >=2 items)
            assert len(doc.tables[2].rows) == 3  # fixed maintenance rows

    elif template_id == "hif-14-key-access":
        assert len(doc.tables) == 3
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[1].rows) >= 3  # access_records (header + >=2 items)
            assert len(doc.tables[2].rows) == 4  # fixed special rules
        # Invariant: special rules rows must not split across pages
        for row in doc.tables[2].rows:
            assert row._tr.trPr is not None and row._tr.trPr.find(docx.oxml.ns.qn("w:cantSplit")) is not None

    elif template_id == "hif-18-purchase":
        assert len(doc.tables) == 6
        assert len(doc.tables[1].rows) == 6  # 6 fixed rows
        assert len(doc.tables[3].rows) == 5  # 5 fixed rows
        assert len(doc.tables[4].rows) == 6  # header + 5 fixed checklist rows
        assert len(doc.tables[5].rows) == 2  # 2 fixed rows
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[2].rows) >= 3  # order_items (header + >=2 items)

    elif template_id == "hif-19-sponsor":
        assert len(doc.tables) == 6
        assert len(doc.tables[1].rows) == 7  # 7 fixed rows
        assert len(doc.tables[2].rows) == 8  # 8 fixed rows
        assert len(doc.tables[3].rows) == 13  # header + 12 fixed checklist rows
        assert len(doc.tables[5].rows) == 4  # 4 fixed rows
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[4].rows) >= 3  # deliverables (header + >=2 items)
        # Invariant: renewal rows must not split across pages
        for row in doc.tables[5].rows:
            assert row._tr.trPr is not None and row._tr.trPr.find(docx.oxml.ns.qn("w:cantSplit")) is not None

    elif template_id == "hif-21-gdpr":
        assert len(doc.tables) == 4
        assert len(doc.tables[3].rows) == 9  # header + 8 fixed checkpoints
        if profile_name in ("normal", "edge"):
            assert len(doc.tables[1].rows) >= 3  # processing_activities (header + >=2 items)
            assert len(doc.tables[2].rows) >= 3  # consent_records (header + >=2 items)
        # Invariant: repeater and control rows must not split across pages, headers repeat
        for table in (doc.tables[1], doc.tables[2], doc.tables[3]):
            assert table.rows[0]._tr.trPr.find(docx.oxml.ns.qn("w:tblHeader")) is not None
            for row in table.rows:
                assert row._tr.trPr is not None and row._tr.trPr.find(docx.oxml.ns.qn("w:cantSplit")) is not None

    # 9. Literal edge content remains literal
    if profile_name == "edge":
        # Danish chars should be preserved
        assert "æ" in doc_text.lower() or "ø" in doc_text.lower() or "å" in doc_text.lower()
        assert "&" in doc_text


@pytest.mark.parametrize("template_id", TEMPLATE_IDS)
def test_pdf_compatibility_real_or_skip(template_id: str, tmp_path: Path) -> None:
    """Real PDF compatibility test for normal profile. Skipped if LibreOffice is unavailable."""
    soffice_cmd = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice_cmd:
        pytest.skip("LibreOffice (soffice/libreoffice) executable is not installed.")

    pkg_dir = TEMPLATES_DIR / template_id
    profile_path = PROFILES_DIR / template_id / "normal.yaml"
    pkg = load_template_package(pkg_dir)

    raw_profile = yaml.safe_load(profile_path.read_text(encoding="utf-8"))
    val_res = validate_all_steps_values(pkg.form_definition, raw_profile)
    assert val_res.is_valid

    context = build_render_context(pkg.form_definition, val_res.coerced_values)
    docx_path = tmp_path / f"{template_id}_normal.docx"
    render_docx(pkg.document_template, context, docx_path)

    pdf_out_dir = tmp_path / "pdf_out"
    pdf_path = convert_docx_to_pdf(docx_path, pdf_out_dir, soffice_cmd=soffice_cmd)
    assert pdf_path.is_file()
    assert pdf_path.stat().st_size > 0
    content = pdf_path.read_bytes()
    assert content.startswith(b"%PDF-")


def test_hif_01_web_journey() -> None:
    """Full web journey for HIF-01 Rollebeskrivelse across all 7 steps."""
    app = create_app(template_root=TEMPLATES_DIR)
    client = TestClient(app)

    # 1. Start journey
    resp = client.post("/templates/hif-01-role/start", follow_redirects=False)
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/0"

    # Step 0: metadata
    resp = client.post(
        "/journey/step/0",
        data={
            "action": "next",
            "owner_committee": "Fodboldudvalget",
            "version_label": "1.0",
            "updated_at": "2026-09-01",
            "approved_by": "Fodboldudvalget / Bestyrelsen",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/1"

    # Step 1: role_info
    resp = client.post(
        "/journey/step/1",
        data={
            "action": "next",
            "role_name": "Ungdomsformand Fodbold",
            "role_area": "Fodbold",
            "role_purpose": "Overordnet ansvar for ungdomsfodbold i Hashøj IF.",
            "reports_to": "Bestyrelsen",
            "collaboration_partners": "Trænere og forældre",
            "estimated_time": "3-5 timer ugentligt",
            "election_period": "2 år",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/2"

    # Step 2: responsibilities_step
    resp = client.post(
        "/journey/step/2",
        data={
            "action": "next",
            "responsibilities.0.task_name": "Holdtilmelding DBU",
            "responsibilities.0.frequency": "Halvårligt",
            "responsibilities.0.deadline": "1. marts",
            "responsibilities.0.relevant_task_card": "HIF-02 DBU Holdkort",
            "responsibilities.1.task_name": "Trænermøder",
            "responsibilities.1.frequency": "Kvartalsvis",
            "responsibilities.1.deadline": "Før sæsonstart",
            "responsibilities.1.relevant_task_card": "HIF-02 Trænermøder",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/3"

    # Step 3: mandate_step
    resp = client.post(
        "/journey/step/3",
        data={
            "action": "next",
            "can_decide_independently": "Må godkende holdtilmeldinger og træningstider.",
            "requires_approval_from": "Dispositioner ud over årsbudgettet.",
            "financial_mandate": "Op til 2.500 kr.",
            "contract_signing_authority": "Nej",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/4"

    # Step 4: resources_step
    resp = client.post(
        "/journey/step/4",
        data={
            "action": "next",
            "access_resources.0.resource_name": "DBU KlubOffice",
            "access_resources.0.purpose": "Spillerregistrering",
            "access_resources.0.access_procedure": "Tildeles af DBU-administrator",
            "access_resources.0.cessation_responsibility": "Lukkes ved fratrædelse",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/5"

    # Step 5: annual_cycle_step
    resp = client.post(
        "/journey/step/5",
        data={
            "action": "next",
            "annual_cycle.0.period": "Marts",
            "annual_cycle.0.activity": "Sæsonopstart",
            "annual_cycle.0.preparation": "Trænermøde",
            "annual_cycle.0.recipient_partner": "Trænere",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/6"

    # Step 6: competencies_and_handover
    resp = client.post(
        "/journey/step/6",
        data={
            "action": "next",
            "required_competencies": "Samarbejdsevner og foreningserfaring.",
            "preferred_competencies": "Kendskab til DBU.",
            "onboarding_description": "Overlevering fra afgående formand.",
            "key_documents": "Vedtægter og politikker.",
            "handover_items": "Koder og nøglebrikker.",
            "documentation_location": "Google Drev / Fodbold",
            "contact_for_questions": "Fodboldformand Lars Nielsen",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/review"

    # Review page
    review_resp = client.get("/journey/review")
    assert review_resp.status_code == 200
    assert "Ungdomsformand Fodbold" in review_resp.text
    assert "Holdtilmelding DBU" in review_resp.text

    # HTML Preview
    preview_resp = client.get("/journey/preview")
    assert preview_resp.status_code == 200
    assert "Ungdomsformand Fodbold" in preview_resp.text

    # DOCX Generation
    gen_resp = client.post("/journey/generate/docx")
    assert gen_resp.status_code == 200
    assert gen_resp.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert len(gen_resp.content) > 0


def test_hif_02_web_journey_with_repeaters() -> None:
    """Full web journey for HIF-02 Opgavekort across all 4 steps with multiple repeaters."""
    app = create_app(template_root=TEMPLATES_DIR)
    client = TestClient(app)

    # 1. Start journey
    resp = client.post("/templates/hif-02-task/start", follow_redirects=False)
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/0"

    # Step 0: basic_info
    resp = client.post(
        "/journey/step/0",
        data={
            "action": "next",
            "owner_committee": "Fodboldudvalget",
            "version_label": "1.0",
            "updated_at": "2026-09-01",
            "approved_by": "Baneudvalget",
            "task_title": "Kridtning af 11-mandsbane",
            "department": "Fodbold",
            "task_type": "Fast",
            "execution_timing": "Fredag eftermiddag før kamp",
            "frequency": "Ugentligt",
            "estimated_duration": "45 minutter",
            "location": "Bane 1 og 2",
            "responsible_role": "Baneformand",
            "purpose": "Sikre tydelige og korrekte banemarkeringer.",
            "completion_criteria": "Alle linjer er kridtet snorlige.",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/1"

    # Step 1: prerequisites_and_materials
    resp = client.post(
        "/journey/step/1",
        data={
            "action": "next",
            "prerequisites.0.prerequisite_item": "Tjek græshøjde",
            "prerequisites.0.location": "Bane 1",
            "prerequisites.0.notes": "Skal være slået",
            "materials.0.material_name": "Kridtvogn",
            "materials.0.quantity_spec": "1 stk.",
            "materials.0.location": "Materielskur",
            "materials.0.notes": "Opladet batteri",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/2"

    # Step 2: execution_steps (task_steps)
    resp = client.post(
        "/journey/step/2",
        data={
            "action": "next",
            "task_steps.0.step_title": "Opbland maling",
            "task_steps.0.step_description": "2L maling og 8L vand",
            "task_steps.0.step_checkpoint": "Ingen klumper",
            "task_steps.0.is_done": "on",
            "task_steps.1.step_title": "Kør linjer op",
            "task_steps.1.step_description": "Følg hjørneflagspunkter",
            "task_steps.1.step_checkpoint": "Linjebredde 10-12 cm",
            "task_steps.1.is_done": "on",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/3"

    # Step 3: safety_and_followup
    resp = client.post(
        "/journey/step/3",
        data={
            "action": "next",
            "safety_risks": "Glat underlag",
            "protective_equipment": "Handsker",
            "qualification_access_requirements": "Instruktion i sprøjtevogn",
            "special_precautions": "Ikke i kraftigt regnvejr",
            "completion_checks.0.checkpoint": "Materielskur aflåst",
            "completion_checks.0.acceptance_criteria": "Hængelås låst",
            "completion_checks.0.is_checked": "on",
            "post_task_documentation": "Notér i driftslog",
            "who_is_informed": "Baneformanden",
            "if_defective_or_missing": "Giv besked hvis maling mangler",
            "deviations_issues": "Ingen afvigelser",
            "improvement_suggestions": "Køb ny si til maling",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/review"

    # Review
    rev_resp = client.get("/journey/review")
    assert rev_resp.status_code == 200
    assert "Kridtning af 11-mandsbane" in rev_resp.text
    assert "Opbland maling" in rev_resp.text

    # HTML Preview
    prev_resp = client.get("/journey/preview")
    assert prev_resp.status_code == 200
    assert "Kridtning af 11-mandsbane" in prev_resp.text

    # DOCX Generation
    docx_resp = client.post("/journey/generate/docx")
    assert docx_resp.status_code == 200
    assert len(docx_resp.content) > 0


def test_hif_07_web_journey_multi_step() -> None:
    """Full web journey for HIF-07 Arrangementsskabelon across all 8 steps."""
    app = create_app(template_root=TEMPLATES_DIR)
    client = TestClient(app)

    # Start
    resp = client.post("/templates/hif-07-event/start", follow_redirects=False)
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/0"

    # Step 0: overview
    resp = client.post(
        "/journey/step/0",
        data={
            "action": "next",
            "owner_committee": "Arrangementsudvalg",
            "version_label": "1.0",
            "updated_at": "2026-09-01",
            "approved_by": "Bestyrelsen",
            "event_title": "Hashøj Sommerfest 2026",
            "date_and_time": "Lørdag d. 13. juni 2026 kl. 10:00 - 22:00",
            "location": "Klubhuset og Baneanlægget",
            "purpose": "Årlig sommerfest til styrkelse af fællesskabet i Hashøj IF.",
            "target_audience": "Alle borgere og medlemmer",
            "expected_participants": "250-300 deltagere",
            "lead_organizer": "Mette Testsen (tlf. 40 50 60 70)",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/1"

    # Step 1: criteria
    resp = client.post(
        "/journey/step/1",
        data={
            "action": "next",
            "success_criteria.0.criterion": "Deltagere",
            "success_criteria.0.measurement_method": "Optælling",
            "success_criteria.0.target_value": "Min. 200",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/2"

    # Step 2: schedule
    resp = client.post(
        "/journey/step/2",
        data={
            "action": "next",
            "timeline.0.time_slot": "1. maj",
            "timeline.0.activity": "Bestille telt",
            "timeline.0.responsible": "Mette",
            "timeline.0.dependency": "Budget godkendt",
            "timeline.0.status": "Udført",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/3"

    # Step 3: staffing_step
    resp = client.post(
        "/journey/step/3",
        data={
            "action": "next",
            "staffing.0.role_title": "Grillmester",
            "staffing.0.count": "2",
            "staffing.0.time_slot": "11:30 - 15:30",
            "staffing.0.assigned_person": "Søren Grillsen",
            "staffing.0.instruction_task_card": "HIF-02 Grill",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/4"

    # Step 4: logistics_step
    resp = client.post(
        "/journey/step/4",
        data={
            "action": "next",
            "logistics.0.item_need": "Festtelt",
            "logistics.0.quantity": "1 stk.",
            "logistics.0.source": "Teltudlejning",
            "logistics.0.responsible": "Mette",
            "logistics.0.ready_by": "Fredag kl. 14",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/5"

    # Step 5: budget_step
    resp = client.post(
        "/journey/step/5",
        data={
            "action": "next",
            "budget_items.0.budget_item": "Salg mad & drikke",
            "budget_items.0.budget_amount": "18.000 kr.",
            "budget_items.0.actual_amount": "19.500 kr.",
            "budget_items.0.comment": "Kiosksalg",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/6"

    # Step 6: comms_step
    resp = client.post(
        "/journey/step/6",
        data={
            "action": "next",
            "communications.0.target_audience": "Medlemmer",
            "communications.0.message": "Invitation til fest",
            "communications.0.channel": "Facebook",
            "communications.0.date_slot": "1. maj",
            "communications.0.responsible": "PR-ansvarlig",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/step/7"

    # Step 7: safety_and_eval
    resp = client.post(
        "/journey/step/7",
        data={
            "action": "next",
            "safety_risks": "Mange gæster og grill",
            "first_aid_evacuation": "Førstehjælpstaske ved indgang",
            "permits_authorities": "Lejlighedstilladelse",
            "emergency_contacts": "Mette Testsen (40 50 60 70)",
            "participant_outcome": "275 deltagere",
            "what_worked": "Grill og hoppeborg",
            "what_to_change_next_time": "Flere skraldespande",
            "financial_result": "+9.200 kr.",
        },
        follow_redirects=False,
    )
    assert resp.status_code == 303
    assert resp.headers["location"] == "/journey/review"

    # Review
    rev_resp = client.get("/journey/review")
    assert rev_resp.status_code == 200
    assert "Hashøj Sommerfest 2026" in rev_resp.text

    # HTML Preview
    prev_resp = client.get("/journey/preview")
    assert prev_resp.status_code == 200
    assert "Hashøj Sommerfest 2026" in prev_resp.text

    # DOCX Generation
    docx_resp = client.post("/journey/generate/docx")
    assert docx_resp.status_code == 200
    assert len(docx_resp.content) > 0


def test_hif_17_incident_activation_gate_disabled() -> None:
    """HIF-17 contains injury/medical and identifiable-person data and must remain disabled

    until a separate human privacy/legal clearance gate is passed.
    """
    pkg = load_template_package(TEMPLATES_DIR / "hif-17-incident")
    assert pkg.form_definition["enabled"] is False

    # Verify that disabled template is omitted from the public index
    app = create_app(template_root=TEMPLATES_DIR)
    client = TestClient(app)
    resp = client.get("/")
    assert resp.status_code == 200
    assert "hif-17-incident" not in resp.text


def test_edge_escaping_authoritative_pipeline_sentinel(tmp_path: Path) -> None:
    """Verify that exact sentinel with special XML/Jinja characters is preserved literally

    in both scalar and repeater child values through the authoritative pipeline.
    """
    sentinel = "A & B < C > D {{ example }}"
    pkg = load_template_package(TEMPLATES_DIR / "hif-01-role")
    fdef = validate_template_definition(pkg.form_definition)
    prof_path = PROFILES_DIR / "hif-01-role" / "normal.yaml"
    raw_prof = yaml.safe_load(prof_path.read_text(encoding="utf-8"))

    raw_prof["role_name"] = sentinel
    raw_prof["responsibilities"][0]["task_name"] = sentinel

    val_res = validate_all_steps_values(fdef, raw_prof)
    assert val_res.is_valid

    context = build_render_context(fdef, val_res.coerced_values)
    out_docx = tmp_path / "sentinel_test.docx"
    render_docx(pkg.document_template, context, out_docx)

    doc = docx.Document(str(out_docx))
    all_text = _extract_all_docx_text(doc)
    assert all_text.count(sentinel) == 2
