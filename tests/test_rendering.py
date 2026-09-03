from pathlib import Path
import docx
import pytest

from hashoej_document_builder.core.errors import DocumentRenderingError
from hashoej_document_builder.core.rendering import build_render_context, render_docx


def _make_template_docx(path: Path, paragraphs: list[str]) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))
    return path


def test_build_render_context_omits_inactive_stale_and_info_fields() -> None:
    form_def = {
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {"id": "role_level", "type": "select", "options": [{"value": "board", "label": "Bestyrelse"}, {"value": "committee", "label": "Udvalg"}]},
                    {"id": "committee_name", "type": "text", "show_when": {"field": "role_level", "equals": "committee"}},
                    {"id": "privacy_info", "type": "info", "text": "Hjælpetekst"},
                    {"id": "optional_notes", "type": "textarea"},
                    {"id": "tasks", "type": "repeater", "fields": [{"id": "title", "type": "text"}, {"id": "info_child", "type": "info", "text": "Hjælp"}]},
                ],
            }
        ]
    }

    # Stale value for committee_name is present in session, but role_level is board (inactive)
    values = {
        "role_level": "board",
        "committee_name": "Stale Udvalgsnavn",
        "optional_notes": None,
        "tasks": [{"title": "Opgave 1"}],
    }

    context = build_render_context(form_def, values)

    assert context["role_level"] == "board"
    assert context["committee_name"] == ""  # Cleared because show_when is inactive
    assert "privacy_info" not in context  # info field never enters render context
    assert context["optional_notes"] == ""
    assert context["tasks"] == [{"title": "Opgave 1"}]
    assert "info_child" not in context["tasks"][0]


def test_render_docx_scalars_danish_chars_and_safe_escaping(tmp_path: Path) -> None:
    tpl_file = _make_template_docx(
        tmp_path / "template.docx",
        [
            "Rolle: {{ role_name }}",
            "Noter: {{ notes }}",
            "JinjaTekst: {{ jinja_text }}",
            "XmlTekst: {{ xml_text }}",
        ],
    )
    out_file = tmp_path / "output.docx"

    context = {
        "role_name": "Formand for Børne- og Ungdomsudvalget (ÆØÅ - æøå)",
        "notes": "A & B < C > D",
        "jinja_text": "{{ 7 * 7 }} og {% if something %}farlig{% endif %}",
        "xml_text": "<script>alert('XSS')</script> & <test>",
    }

    result_path = render_docx(tpl_file, context, out_file)
    assert result_path.is_file()

    # Inspect rendered docx using python-docx
    doc = docx.Document(str(result_path))
    texts = [p.text for p in doc.paragraphs]

    assert texts[0] == "Rolle: Formand for Børne- og Ungdomsudvalget (ÆØÅ - æøå)"
    assert texts[1] == "Noter: A & B < C > D"
    assert texts[2] == "JinjaTekst: {{ 7 * 7 }} og {% if something %}farlig{% endif %}"
    assert texts[3] == "XmlTekst: <script>alert('XSS')</script> & <test>"


def test_render_docx_repeater_multiple_rows_and_empty(tmp_path: Path) -> None:
    tpl_file = _make_template_docx(
        tmp_path / "template.docx",
        [
            "Opgaver start:",
            "{% for item in tasks %}",
            "Opgave: {{ item.title }} - Timer: {{ item.hours }}",
            "{% endfor %}",
            "Opgaver slut.",
        ],
    )

    # 1. Multiple rows
    out_file1 = tmp_path / "output1.docx"
    render_docx(
        tpl_file,
        {"tasks": [{"title": "Regnskab", "hours": 10}, {"title": "Møder", "hours": 5}]},
        out_file1,
    )
    doc1 = docx.Document(str(out_file1))
    texts1 = [p.text for p in doc1.paragraphs if p.text.strip()]
    assert "Opgave: Regnskab - Timer: 10" in texts1
    assert "Opgave: Møder - Timer: 5" in texts1

    # 2. Empty repeater -> zero loop iterations
    out_file2 = tmp_path / "output2.docx"
    render_docx(tpl_file, {"tasks": []}, out_file2)
    doc2 = docx.Document(str(out_file2))
    texts2 = [p.text for p in doc2.paragraphs if p.text.strip()]
    assert texts2 == ["Opgaver start:", "Opgaver slut."]


def test_render_docx_atomic_failure_cleanup(tmp_path: Path) -> None:
    non_existent_tpl = tmp_path / "missing.docx"
    out_file = tmp_path / "output_fail.docx"

    with pytest.raises(DocumentRenderingError, match="DOCX template file not found"):
        render_docx(non_existent_tpl, {}, out_file)

    assert not out_file.exists()
