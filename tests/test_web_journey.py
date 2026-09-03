from datetime import datetime, timedelta, timezone
import logging
from pathlib import Path
import pytest
from starlette.testclient import TestClient
import yaml

import docx
from hashoej_document_builder.core.session import SessionStore
from hashoej_document_builder.web.app import COOKIE_NAME, DOCX_MIME_TYPE, PDF_MIME_TYPE, create_app


def _create_synthetic_test_template(
    dir_path: Path,
    template_id: str = "hif-01-role",
    version: int = 1,
    enabled: bool = True,
) -> Path:
    dir_path.mkdir(parents=True, exist_ok=True)
    raw_dict = {
        "id": template_id,
        "version": version,
        "enabled": enabled,
        "title": "Rollebeskrivelse" if template_id == "hif-01-role" else "Opgavekort",
        "category": "Organisation",
        "description": "Beskriv en rolle i Hashøj IF.",
        "steps": [
            {
                "id": "step1",
                "title": "Grundoplysninger",
                "description": "Indtast rollens basisoplysninger",
                "fields": [
                    {
                        "id": "role_name",
                        "type": "text",
                        "label": "Rollenavn",
                        "required": True,
                        "default": "Formand",
                    },
                    {
                        "id": "role_level",
                        "type": "select",
                        "label": "Niveau",
                        "options": [
                            {"value": "board", "label": "Bestyrelse"},
                            {"value": "committee", "label": "Udvalg"},
                        ],
                        "default": "board",
                    },
                    {
                        "id": "committee_name",
                        "type": "text",
                        "label": "Udvalg",
                        "required": True,
                        "show_when": {"field": "role_level", "equals": "committee"},
                    },
                    {
                        "id": "has_key",
                        "type": "checkbox",
                        "label": "Har nøgle",
                    },
                    {
                        "id": "sports",
                        "type": "multiselect",
                        "label": "Sportsgrene",
                        "options": [
                            {"value": "f", "label": "Fodbold"},
                            {"value": "g", "label": "Gymnastik"},
                        ],
                    },
                ],
            },
            {
                "id": "step2",
                "title": "Opgaver og samtykke",
                "fields": [
                    {
                        "id": "tasks",
                        "type": "repeater",
                        "label": "Opgaver",
                        "min_items": 1,
                        "max_items": 3,
                        "fields": [
                            {"id": "title", "type": "text", "label": "Opgavetitel", "required": True, "default": "Standardopgave"},
                            {"id": "frequency", "type": "text", "label": "Hyppighed", "required": False},
                            {"id": "priority", "type": "number", "label": "Prioritet", "required": False, "default": 1},
                            {
                                "id": "task_type",
                                "type": "select",
                                "label": "Type",
                                "options": [
                                    {"value": "admin", "label": "Administration"},
                                    {"value": "sport", "label": "Sportsligt"},
                                ],
                                "default": "admin",
                            },
                            {
                                "id": "child_info",
                                "type": "info",
                                "text": "Husk at angive realistiske tidsestimater.",
                            },
                        ],
                    },
                    {
                        "id": "privacy_info",
                        "type": "info",
                        "variant": "privacy",
                        "text": "Data gemmes kun midlertidigt i serverens RAM.",
                    },
                    {
                        "id": "ack",
                        "type": "checkbox",
                        "label": "Jeg forstår privatlivsinformationen",
                        "purpose": "acknowledgement",
                        "required": True,
                    },
                ],
            },
        ],
    }
    (dir_path / "template.yaml").write_text(yaml.dump(raw_dict), encoding="utf-8")

    doc = docx.Document()
    doc.add_paragraph("Rolle: {{ role_name }} - Niveau: {{ role_level }}")
    doc.add_paragraph("{% for item in tasks %}{{ item.title }} - {{ item.frequency }}{% endfor %}")
    doc.save(str(dir_path / "document.docx"))

    return dir_path


@pytest.fixture
def app_and_store(tmp_path: Path):
    template_root = tmp_path / "templates"
    template_root.mkdir()
    _create_synthetic_test_template(template_root / "hif-01-role", template_id="hif-01-role", enabled=True)
    _create_synthetic_test_template(template_root / "hif-02-task", template_id="hif-02-task", enabled=False)

    artifacts_root = tmp_path / "artifacts"
    from hashoej_document_builder.core.artifacts import ArtifactManager
    art_mgr = ArtifactManager(temp_root=artifacts_root)

    store = SessionStore()
    app = create_app(template_root=template_root, session_store=store, artifact_manager=art_mgr)
    return app, store, template_root


def test_homepage_shows_enabled_and_hides_disabled_templates(app_and_store):
    app, _, _ = app_and_store
    client = TestClient(app)

    response = client.get("/")
    assert response.status_code == 200
    html = response.text

    assert "Rollebeskrivelse" in html
    assert "/templates/hif-01-role" in html
    assert "hif-02-task" not in html


def test_homepage_empty_state(tmp_path: Path):
    empty_root = tmp_path / "empty_templates"
    empty_root.mkdir()
    app = create_app(template_root=empty_root)
    client = TestClient(app)

    response = client.get("/")
    assert response.status_code == 200
    assert "Ingen skabeloner tilgængelige" in response.text


def test_intro_page_for_enabled_and_disabled_templates(app_and_store):
    app, _, _ = app_and_store
    client = TestClient(app)

    # Enabled template intro
    response = client.get("/templates/hif-01-role")
    assert response.status_code == 200
    assert "Opret rollebeskrivelse" in response.text
    assert "Sådan fungerer det" in response.text

    # Disabled template intro is 404
    resp_disabled = client.get("/templates/hif-02-task")
    assert resp_disabled.status_code == 404
    assert "Dokumenttype ikke fundet" in resp_disabled.text

    # Unknown template is 404
    resp_unknown = client.get("/templates/non-existent-template")
    assert resp_unknown.status_code == 404


def test_start_journey_sets_secure_opaque_cookie_and_redirects(app_and_store):
    app, store, _ = app_and_store
    client = TestClient(app)

    response = client.post("/templates/hif-01-role/start", follow_redirects=False)
    assert response.status_code == 303
    assert response.headers["location"] == "/journey/step/0"

    assert COOKIE_NAME == "document_builder_session_id"
    assert COOKIE_NAME in response.cookies
    cookie_val = response.cookies[COOKIE_NAME]
    assert len(cookie_val) >= 32

    # Verify cookie does NOT contain form data or template definitions
    assert "Formand" not in cookie_val
    assert "hif-01-role" not in cookie_val

    session = store.get(cookie_val)
    assert session is not None
    assert session.template_id == "hif-01-role"
    assert session.current_step == 0
    assert session.values["role_name"] == "Formand"


def test_cookie_https_and_proxy_trust(tmp_path: Path):
    template_root = tmp_path / "templates"
    template_root.mkdir()
    _create_synthetic_test_template(template_root / "hif-01-role", template_id="hif-01-role", enabled=True)

    app = create_app(template_root=template_root)

    # 1. Real HTTPS request -> Secure cookie
    client_https = TestClient(app, base_url="https://testserver")
    resp_https = client_https.post("/templates/hif-01-role/start", follow_redirects=False)
    set_cookie_https = resp_https.headers.get("set-cookie", "")
    assert "Secure" in set_cookie_https
    assert "HttpOnly" in set_cookie_https
    assert "samesite=lax" in set_cookie_https.lower()

    # 2. Plain HTTP request -> No Secure cookie
    client_http = TestClient(app, base_url="http://testserver")
    resp_http = client_http.post("/templates/hif-01-role/start", follow_redirects=False)
    set_cookie_http = resp_http.headers.get("set-cookie", "")
    assert "Secure" not in set_cookie_http

    # 3. Plain HTTP with spoofed X-Forwarded-Proto -> Must not be trusted alone
    resp_spoof = client_http.post(
        "/templates/hif-01-role/start",
        headers={"X-Forwarded-Proto": "https"},
        follow_redirects=False,
    )
    set_cookie_spoof = resp_spoof.headers.get("set-cookie", "")
    assert "Secure" not in set_cookie_spoof


def test_progression_boundary_and_url_tampering(app_and_store):
    app, store, _ = app_and_store
    client = TestClient(app)

    # Start session: current_step is 0
    client.post("/templates/hif-01-role/start")

    # Direct GET to locked Step 1 must redirect to highest allowed (Step 0)
    resp_jump = client.get("/journey/step/1", follow_redirects=False)
    assert resp_jump.status_code == 303
    assert resp_jump.headers["location"] == "/journey/step/0"

    # Direct POST to locked Step 1 must redirect to highest allowed (Step 0)
    resp_post_jump = client.post(
        "/journey/step/1",
        data={"tasks.0.title": "Illegal", "ack": "true", "action": "next"},
        follow_redirects=False,
    )
    assert resp_post_jump.status_code == 303
    assert resp_post_jump.headers["location"] == "/journey/step/0"

    # Direct GET to Review must redirect to highest allowed (Step 0)
    resp_rev_jump = client.get("/journey/review", follow_redirects=False)
    assert resp_rev_jump.status_code == 303
    assert resp_rev_jump.headers["location"] == "/journey/step/0"

    # Complete Step 0 successfully -> unlocks Step 1 (current_step becomes 1)
    step0_ok = client.post(
        "/journey/step/0",
        data={"role_name": "Kasserer", "role_level": "opt_0", "action": "next"},
        follow_redirects=False,
    )
    assert step0_ok.status_code == 303
    assert step0_ok.headers["location"] == "/journey/step/1"

    # Now GET step 1 is allowed
    assert client.get("/journey/step/1").status_code == 200

    # Back button to Step 0 does NOT reduce current_step boundary
    back_resp = client.post("/journey/step/1", data={"action": "prev"}, follow_redirects=False)
    assert back_resp.status_code == 303
    assert back_resp.headers["location"] == "/journey/step/0"

    # From Step 0, Step 1 is still unlocked
    assert client.get("/journey/step/1").status_code == 200


def test_complete_guided_journey_flow(app_and_store):
    app, store, _ = app_and_store
    client = TestClient(app)

    # 1. Start journey
    start_resp = client.post("/templates/hif-01-role/start", follow_redirects=True)
    assert start_resp.status_code == 200
    assert "Trin 1 af 2" in start_resp.text

    # 2. Validation error on Step 1: clear role_name and submit
    err_resp = client.post(
        "/journey/step/0",
        data={"role_name": "", "role_level": "opt_0", "action": "next"},
    )
    assert err_resp.status_code == 422
    assert "Dette felt er påkrævet." in err_resp.text

    # 3. Refreshing Step 0 preserves typed values via draft
    ref_resp = client.get("/journey/step/0")
    assert ref_resp.status_code == 200

    # 4. Valid submission on Step 1: select 'board' (committee_name is inactive)
    step1_resp = client.post(
        "/journey/step/0",
        data={"role_name": "Kasserer", "role_level": "opt_0", "has_key": "true", "sports": ["opt_0"], "action": "next"},
        follow_redirects=True,
    )
    assert step1_resp.status_code == 200
    assert "Trin 2 af 2" in step1_resp.text

    # 5. Backward navigation: preserves entered values without modifying authoritative values
    back_resp = client.post(
        "/journey/step/1",
        data={"action": "prev"},
        follow_redirects=True,
    )
    assert back_resp.status_code == 200
    assert "Kasserer" in back_resp.text

    # Advance to Step 2 again
    client.post(
        "/journey/step/0",
        data={"role_name": "Kasserer", "role_level": "opt_0", "has_key": "true", "sports": ["opt_0"], "action": "next"},
        follow_redirects=True,
    )

    # 6. Repeater operations on Step 2: Add row with defaults
    add_row_resp = client.post(
        "/journey/step/1",
        data={
            "action": "repeater_add:tasks",
        },
        follow_redirects=True,
    )
    assert add_row_resp.status_code == 200
    assert "Standardopgave" in add_row_resp.text

    # 7. Valid Step 2 submission -> redirects to Review
    step2_ok = client.post(
        "/journey/step/1",
        data={
            "tasks.0.title": "Regnskab",
            "tasks.0.frequency": "Månedligt",
            "tasks.0.task_type": "opt_0",
            "ack": "true",
            "action": "next",
        },
        follow_redirects=True,
    )
    assert step2_ok.status_code == 200
    assert "Kontrollér dine svar" in step2_ok.text
    assert "Kasserer" in step2_ok.text
    assert "Regnskab" in step2_ok.text
    assert "Bestyrelse" in step2_ok.text

    # 8. HTML Preview
    preview_resp = client.get("/journey/preview")
    assert preview_resp.status_code == 200
    assert "Sådan kommer dokumentet til at se ud" in preview_resp.text
    assert "Kasserer" in preview_resp.text
    assert "Regnskab" in preview_resp.text
    assert "Bestyrelse" in preview_resp.text
    assert "Data gemmes kun midlertidigt" not in preview_resp.text


def test_empty_controls_preserved_in_draft_after_validation_error(app_and_store):
    """Revisiting a step, unchecking a checkbox/clearing multiselect, and failing validation preserves empty state."""
    app, store, _ = app_and_store
    client = TestClient(app)

    start_resp = client.post("/templates/hif-01-role/start", follow_redirects=False)
    cookie_val = start_resp.cookies[COOKIE_NAME]

    # Complete Step 0 with checkbox=True and sports=['f', 'g']
    client.post(
        "/journey/step/0",
        data={"role_name": "Kasserer", "role_level": "opt_0", "has_key": "true", "sports": ["opt_0", "opt_1"], "action": "next"},
    )

    session_before = store.get(cookie_val)
    assert session_before.values["has_key"] is True
    assert session_before.values["sports"] == ["f", "g"]

    # Revisit Step 0: uncheck checkbox, clear sports, and clear role_name to force validation error
    err_resp = client.post(
        "/journey/step/0",
        data={"role_name": "", "role_level": "opt_0", "action": "next"},
    )
    assert err_resp.status_code == 422

    # Refresh step 0: draft must show checkbox unchecked and multiselect unselected
    refresh_resp = client.get("/journey/step/0")
    assert refresh_resp.status_code == 200
    html = refresh_resp.text

    # Checkbox 'has_key' must NOT be checked
    assert 'name="has_key" value="true"\n             checked' not in html

    # Multiselect 'sports' options must NOT be checked
    assert 'name="sports" value="opt_0"\n                 checked' not in html


def test_repeater_child_info_never_in_session_or_preview_table(app_and_store):
    app, store, _ = app_and_store
    client = TestClient(app)

    start_resp = client.post("/templates/hif-01-role/start", follow_redirects=False)
    cookie_val = start_resp.cookies[COOKIE_NAME]

    client.post(
        "/journey/step/0",
        data={"role_name": "Kasserer", "role_level": "opt_0", "action": "next"},
    )
    client.post(
        "/journey/step/1",
        data={
            "tasks.0.title": "Regnskab",
            "ack": "true",
            "action": "next",
        },
    )

    session = store.get(cookie_val)
    assert session is not None
    # Child info must never be a domain key in session values
    assert "child_info" not in session.values["tasks"][0]

    # Review and Preview must not contain a column header for the info child
    review_resp = client.get("/journey/review")
    assert "<th>Husk at angive" not in review_resp.text

    preview_resp = client.get("/journey/preview")
    assert "<th>Husk at angive" not in preview_resp.text


def test_action_parsing_hardening_tampering(app_and_store):
    app, store, _ = app_and_store
    client = TestClient(app)

    start_resp = client.post("/templates/hif-01-role/start", follow_redirects=False)
    cookie_val = start_resp.cookies[COOKIE_NAME]

    # 1. Duplicate action fields
    dup_resp = client.post(
        "/journey/step/0",
        data={"role_name": "Kasserer", "action": ["next", "prev"]},
        follow_redirects=False,
    )
    assert dup_resp.status_code == 303
    assert dup_resp.headers["location"] == "/journey/step/0"
    # Progression was NOT advanced
    assert store.get(cookie_val).current_step == 0

    # 2. Unknown action string
    unknown_resp = client.post(
        "/journey/step/0",
        data={"role_name": "Kasserer", "action": "delete_all"},
        follow_redirects=False,
    )
    assert unknown_resp.status_code == 303
    assert unknown_resp.headers["location"] == "/journey/step/0"
    assert store.get(cookie_val).current_step == 0

    # 3. Malformed repeater action
    bad_rep = client.post(
        "/journey/step/0",
        data={"action": "repeater_add:"},
        follow_redirects=False,
    )
    assert bad_rep.status_code == 303


def test_option_tokens_roundtrip_scalar_types(tmp_path: Path):
    template_root = tmp_path / "templates"
    template_root.mkdir()
    dir_path = template_root / "hif-test-types"
    dir_path.mkdir(parents=True)
    raw_dict = {
        "id": "hif-test-types",
        "version": 1,
        "enabled": True,
        "title": "Type Test",
        "category": "Test",
        "description": "Test scalar types",
        "steps": [
            {
                "id": "step1",
                "title": "Type Step",
                "fields": [
                    {
                        "id": "opt_scalar",
                        "type": "select",
                        "label": "Scalar Valg",
                        "options": [
                            {"value": 1, "label": "Tal 1"},
                            {"value": "1", "label": "Tekst 1"},
                            {"value": True, "label": "Sand"},
                            {"value": False, "label": "Falsk"},
                        ],
                    },
                    {
                        "id": "cond_field",
                        "type": "text",
                        "label": "Kun Sand",
                        "required": True,
                        "show_when": {"field": "opt_scalar", "equals": True},
                    },
                ],
            },
        ],
    }
    (dir_path / "template.yaml").write_text(yaml.dump(raw_dict), encoding="utf-8")
    doc = docx.Document()
    doc.add_paragraph("Valg: {{ opt_scalar }}")
    doc.save(str(dir_path / "document.docx"))

    store = SessionStore()
    app = create_app(template_root=template_root, session_store=store)
    client = TestClient(app)

    start_resp = client.post("/templates/hif-test-types/start", follow_redirects=False)
    cookie_val = start_resp.cookies[COOKIE_NAME]

    # 1. Select opt_0 -> integer 1 (cond_field is inactive)
    client.post("/journey/step/0", data={"opt_scalar": "opt_0", "action": "next"})
    val_int = store.get(cookie_val).values["opt_scalar"]
    assert val_int == 1
    assert isinstance(val_int, int) and not isinstance(val_int, bool)

    # 2. Select opt_1 -> string "1" (cond_field is inactive)
    client.post("/journey/step/0", data={"opt_scalar": "opt_1", "action": "next"})
    val_str = store.get(cookie_val).values["opt_scalar"]
    assert val_str == "1"
    assert isinstance(val_str, str)

    # 3. Select opt_3 -> boolean False (cond_field is inactive)
    client.post("/journey/step/0", data={"opt_scalar": "opt_3", "action": "next"})
    val_false = store.get(cookie_val).values["opt_scalar"]
    assert val_false is False
    assert isinstance(val_false, bool)

    # 4. Select opt_2 -> boolean True (cond_field IS active and required)
    err_resp = client.post("/journey/step/0", data={"opt_scalar": "opt_2", "cond_field": "", "action": "next"})
    assert err_resp.status_code == 422
    assert "Dette felt er påkrævet." in err_resp.text


def test_field_injection_protection(app_and_store):
    """Unknown or malicious submitted fields must never enter authoritative values or drafts."""
    app, store, _ = app_and_store
    client = TestClient(app)

    start_resp = client.post("/templates/hif-01-role/start", follow_redirects=False)
    cookie_val = start_resp.cookies[COOKIE_NAME]

    # Post malicious injected fields on Step 0
    client.post(
        "/journey/step/0",
        data={
            "role_name": "Kasserer",
            "role_level": "opt_0",
            "action": "next",
            "admin": "true",
            "__internal": "injected",
            "unknown_field": "exploit",
        },
    )

    session = store.get(cookie_val)
    assert session is not None
    assert "admin" not in session.values
    assert "__internal" not in session.values
    assert "unknown_field" not in session.values
    assert "action" not in session.values
    assert session.draft_values == {}


def test_repeater_tampered_actions_and_limits(app_and_store):
    app, store, _ = app_and_store
    client = TestClient(app)

    client.post("/templates/hif-01-role/start")
    client.post(
        "/journey/step/0",
        data={"role_name": "Kasserer", "role_level": "opt_0", "action": "next"},
    )

    # 1. Action on unknown field -> ignored safely without 500 error
    resp_unknown_add = client.post(
        "/journey/step/1",
        data={"action": "repeater_add:unknown_repeater_field"},
        follow_redirects=False,
    )
    assert resp_unknown_add.status_code == 303

    # 2. Malformed remove action -> ignored safely without 500 error
    resp_bad_remove = client.post(
        "/journey/step/1",
        data={"action": "repeater_remove:tasks:not_an_int"},
        follow_redirects=False,
    )
    assert resp_bad_remove.status_code == 303

    # 3. Remove with out-of-range index -> ignored safely
    resp_oob_remove = client.post(
        "/journey/step/1",
        data={"action": "repeater_remove:tasks:999"},
        follow_redirects=False,
    )
    assert resp_oob_remove.status_code == 303

    # 4. Add rows up to max_items (max_items is 3)
    client.post("/journey/step/1", data={"action": "repeater_add:tasks"})
    client.post("/journey/step/1", data={"action": "repeater_add:tasks"})
    client.post("/journey/step/1", data={"action": "repeater_add:tasks"})
    # Fourth add attempt exceeds max_items=3 and is prevented
    client.post("/journey/step/1", data={"action": "repeater_add:tasks"})

    step1_page = client.get("/journey/step/1")
    assert step1_page.status_code == 200
    # Must have exactly 3 rows
    assert 'name="tasks.2._row"' in step1_page.text
    assert 'name="tasks.3._row"' not in step1_page.text


def test_human_readable_option_labels_in_review_and_preview(app_and_store):
    app, store, _ = app_and_store
    client = TestClient(app)

    client.post("/templates/hif-01-role/start")
    client.post(
        "/journey/step/0",
        data={"role_name": "Kasserer", "role_level": "opt_0", "action": "next"},
    )
    client.post(
        "/journey/step/1",
        data={
            "tasks.0.title": "Regnskab",
            "tasks.0.task_type": "opt_0",
            "ack": "true",
            "action": "next",
        },
    )

    # Review page check
    review_resp = client.get("/journey/review")
    assert review_resp.status_code == 200
    # Option value 'board' must render as label 'Bestyrelse'
    assert "Bestyrelse" in review_resp.text
    # Child option value 'admin' must render as label 'Administration'
    assert "Administration" in review_resp.text

    # Preview page check
    preview_resp = client.get("/journey/preview")
    assert preview_resp.status_code == 200
    assert "Bestyrelse" in preview_resp.text
    assert "Administration" in preview_resp.text


def test_session_bound_to_template_id_and_version(tmp_path: Path):
    template_root = tmp_path / "templates"
    template_root.mkdir()
    tpl_dir = _create_synthetic_test_template(
        template_root / "hif-01-role",
        template_id="hif-01-role",
        version=1,
        enabled=True,
    )

    store = SessionStore()
    app = create_app(template_root=template_root, session_store=store)
    client = TestClient(app)

    # Start session bound to version 1
    start_resp = client.post("/templates/hif-01-role/start", follow_redirects=False)
    cookie_val = start_resp.cookies[COOKIE_NAME]
    assert store.get(cookie_val) is not None

    # Step 0 works against version 1
    assert client.get("/journey/step/0").status_code == 200

    # Upgrade installed template on disk to version 2
    _create_synthetic_test_template(
        tpl_dir,
        template_id="hif-01-role",
        version=2,
        enabled=True,
    )

    # Subsequent request with session bound to v1 must fail safely and purge session
    mismatch_resp = client.get("/journey/step/0")
    assert mismatch_resp.status_code == 404
    assert "Skabelonen er opdateret eller ikke længere tilgængelig" in mismatch_resp.text

    # Session deleted from store
    assert store.get(cookie_val) is None


def test_clock_injected_inactivity_ttl(tmp_path: Path):
    template_root = tmp_path / "templates"
    template_root.mkdir()
    _create_synthetic_test_template(template_root / "hif-01-role", template_id="hif-01-role", enabled=True)

    current_time = datetime(2026, 9, 1, 12, 0, 0, tzinfo=timezone.utc)

    def clock():
        return current_time

    app = create_app(template_root=template_root, clock=clock)
    client = TestClient(app)

    client.post("/templates/hif-01-role/start")

    # 30 minutes later: active
    current_time += timedelta(minutes=30)
    assert client.get("/journey/step/0").status_code == 200

    # 45 minutes after last activity (75 min total): active
    current_time += timedelta(minutes=45)
    assert client.get("/journey/step/0").status_code == 200

    # 61 minutes after last activity without touching: expired!
    current_time += timedelta(minutes=61)
    exp_resp = client.get("/journey/step/0")
    assert exp_resp.status_code == 200
    assert "Session udløbet" in exp_resp.text


def test_html_preview_escapes_xss_and_code_like_defaults(app_and_store):
    app, store, _ = app_and_store
    client = TestClient(app)

    client.post("/templates/hif-01-role/start")

    xss_payload = "<script>alert('pwned')</script>"
    jinja_payload = "{{ 7 * 7 }}"
    special_chars = "A & B < C > D"

    client.post(
        "/journey/step/0",
        data={"role_name": xss_payload, "role_level": "opt_0", "action": "next"},
    )
    client.post(
        "/journey/step/1",
        data={
            "tasks.0.title": jinja_payload,
            "tasks.0.frequency": special_chars,
            "ack": "true",
            "action": "next",
        },
    )

    preview = client.get("/journey/preview")
    assert preview.status_code == 200
    text = preview.text

    assert "<script>alert('pwned')</script>" not in text
    assert "&lt;script&gt;alert(&#39;pwned&#39;)&lt;/script&gt;" in text or "&lt;script&gt;alert('pwned')&lt;/script&gt;" in text
    assert "49" not in text
    assert "{{ 7 * 7 }}" in text
    assert "A &amp; B &lt; C &gt; D" in text


def test_privacy_and_logging_discipline(app_and_store, caplog):
    app, store, _ = app_and_store
    client = TestClient(app)

    client.post("/templates/hif-01-role/start")

    secret_value = "SUPER_SECRET_PAYLOAD_ABC123"
    with caplog.at_level(logging.DEBUG):
        resp = client.post(
            "/journey/step/0",
            data={"role_name": "", "committee_name": secret_value, "action": "next"},
        )
        assert resp.status_code == 422

    for record in caplog.records:
        assert secret_value not in record.message
        assert secret_value not in str(record.args)


def test_no_client_side_storage_or_htmx_history_caching(app_and_store):
    app, _, _ = app_and_store
    client = TestClient(app)

    client.post("/templates/hif-01-role/start")
    step_resp = client.get("/journey/step/0")
    assert step_resp.status_code == 200
    html = step_resp.text

    assert 'hx-history="false"' in html
    assert "localStorage" not in html
    assert "sessionStorage" not in html


def test_generate_docx_success_and_mime_headers(app_and_store, tmp_path):
    app, store, _ = app_and_store
    client = TestClient(app)
    artifacts_root = tmp_path / "artifacts"

    start_resp = client.post("/templates/hif-01-role/start", follow_redirects=False)
    cookie_val = start_resp.cookies[COOKIE_NAME]

    # Complete Step 0 and Step 1
    client.post("/journey/step/0", data={"role_name": "Kasserer", "role_level": "opt_0", "action": "next"})
    client.post("/journey/step/1", data={"tasks.0.title": "Regnskab", "ack": "true", "action": "next"})

    # Generation request
    resp = client.post("/journey/generate/docx")
    assert resp.status_code == 200
    assert resp.headers["content-type"] == DOCX_MIME_TYPE
    assert 'filename="hif-01-role.docx"' in resp.headers.get("content-disposition", "")

    # Content is a valid Word document
    import io
    docx_doc = docx.Document(io.BytesIO(resp.content))
    paragraphs_text = " ".join(p.text for p in docx_doc.paragraphs)
    assert "Kasserer" in paragraphs_text
    assert "Regnskab" in paragraphs_text

    # After FileResponse completes, dedicated artifact directory has been deleted
    assert len(list(artifacts_root.iterdir())) == 0

    # Session remains active and unmodified
    session_after = store.get(cookie_val)
    assert session_after is not None
    assert session_after.values["role_name"] == "Kasserer"


def test_generate_pdf_success_and_mime_headers(app_and_store, monkeypatch, tmp_path):
    app, store, _ = app_and_store
    client = TestClient(app)
    artifacts_root = tmp_path / "artifacts"

    start_resp = client.post("/templates/hif-01-role/start", follow_redirects=False)
    cookie_val = start_resp.cookies[COOKIE_NAME]

    client.post("/journey/step/0", data={"role_name": "Kasserer", "role_level": "opt_0", "action": "next"})
    client.post("/journey/step/1", data={"tasks.0.title": "Regnskab", "ack": "true", "action": "next"})

    # Mock convert_docx_to_pdf to simulate successful PDF generation
    def mock_convert(docx_path, output_dir, **kwargs):
        pdf_file = Path(output_dir) / f"{Path(docx_path).stem}.pdf"
        pdf_file.write_bytes(b"%PDF-1.4 mock pdf content")
        return pdf_file

    monkeypatch.setattr("hashoej_document_builder.core.pdf.convert_docx_to_pdf", mock_convert)

    resp = client.post("/journey/generate/pdf")
    assert resp.status_code == 200
    assert resp.headers["content-type"] == PDF_MIME_TYPE
    assert 'filename="hif-01-role.pdf"' in resp.headers.get("content-disposition", "")
    assert resp.content.startswith(b"%PDF-1.4")

    # After FileResponse completes, dedicated directory containing both DOCX and PDF has been deleted
    assert len(list(artifacts_root.iterdir())) == 0

    # Session survives PDF generation
    session_after = store.get(cookie_val)
    assert session_after is not None


def test_generate_pdf_failure_independence_leaves_docx_usable(app_and_store, monkeypatch, tmp_path):
    """If PDF conversion fails, DOCX remains fully functional and downloadable, and temp files are cleaned."""
    app, store, _ = app_and_store
    client = TestClient(app)
    artifacts_root = tmp_path / "artifacts"

    start_resp = client.post("/templates/hif-01-role/start", follow_redirects=False)
    cookie_val = start_resp.cookies[COOKIE_NAME]

    client.post("/journey/step/0", data={"role_name": "Kasserer", "role_level": "opt_0", "action": "next"})
    client.post("/journey/step/1", data={"tasks.0.title": "Regnskab", "ack": "true", "action": "next"})

    # Mock PDF failure
    from hashoej_document_builder.core.errors import PDFConversionUnavailableError
    def mock_fail_convert(*args, **kwargs):
        raise PDFConversionUnavailableError("LibreOffice is not installed.")

    monkeypatch.setattr("hashoej_document_builder.core.pdf.convert_docx_to_pdf", mock_fail_convert)

    pdf_resp = client.post("/journey/generate/pdf")
    # Returns safe error page informing the user and preserving session
    assert pdf_resp.status_code == 200
    assert "PDF-konvertering ikke tilgængelig" in pdf_resp.text
    assert "Hent DOCX" in pdf_resp.text
    assert '<form method="post" action="/journey/generate/docx"' in pdf_resp.text
    assert '<a href="/journey/preview"' in pdf_resp.text

    # Temporary artifact directory for the failed PDF attempt has been cleaned up
    assert len(list(artifacts_root.iterdir())) == 0

    # DOCX generation STILL works cleanly
    docx_resp = client.post("/journey/generate/docx")
    assert docx_resp.status_code == 200
    assert docx_resp.headers["content-type"] == DOCX_MIME_TYPE

    # DOCX artifact directory cleaned up after response
    assert len(list(artifacts_root.iterdir())) == 0

    # Session is preserved
    assert store.get(cookie_val) is not None


def test_generate_docx_rendering_failure_cleans_artifact_dir(app_and_store, monkeypatch, tmp_path):
    app, store, _ = app_and_store
    client = TestClient(app)
    artifacts_root = tmp_path / "artifacts"

    client.post("/templates/hif-01-role/start")
    client.post("/journey/step/0", data={"role_name": "Kasserer", "role_level": "opt_0", "action": "next"})
    client.post("/journey/step/1", data={"tasks.0.title": "Regnskab", "ack": "true", "action": "next"})

    from hashoej_document_builder.core.errors import DocumentRenderingError
    def mock_fail_render(*args, **kwargs):
        raise DocumentRenderingError("DOCX engine broke.")

    monkeypatch.setattr("hashoej_document_builder.core.rendering.render_docx", mock_fail_render)

    resp = client.post("/journey/generate/docx")
    assert resp.status_code == 500
    assert "Dokumentgenerering fejlede" in resp.text

    # Dedicated artifact directory was removed on failure
    assert len(list(artifacts_root.iterdir())) == 0


def test_generate_endpoints_enforce_progression_and_validation(app_and_store):
    app, store, _ = app_and_store
    client = TestClient(app)

    # 1. Start journey: current_step is 0 (wizard incomplete)
    client.post("/templates/hif-01-role/start")

    # Direct generate before wizard completion redirects to current step
    docx_jump = client.post("/journey/generate/docx", follow_redirects=False)
    assert docx_jump.status_code == 303
    assert docx_jump.headers["location"] == "/journey/step/0"

    pdf_jump = client.post("/journey/generate/pdf", follow_redirects=False)
    assert pdf_jump.status_code == 303
    assert pdf_jump.headers["location"] == "/journey/step/0"
