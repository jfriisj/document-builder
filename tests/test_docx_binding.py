from pathlib import Path
import docx
import pytest

from hashoej_document_builder.core.discovery import load_template_package
from hashoej_document_builder.core.docx_binding import validate_docx_binding
from hashoej_document_builder.core.errors import DOCXBindingValidationError
from hashoej_document_builder.core.rendering import render_docx


def _make_docx(path: Path, paragraphs: list[str]) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))
    return path


def test_valid_scalar_binding(tmp_path: Path) -> None:
    docx_file = _make_docx(tmp_path / "document.docx", ["Rolle: {{ role_name }}"])
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [{"id": "role_name", "type": "text", "label": "Rollenavn"}],
            }
        ],
    }
    # Must pass without raising
    validate_docx_binding(docx_file, form_def)


def test_valid_repeater_binding(tmp_path: Path) -> None:
    docx_file = _make_docx(
        tmp_path / "document.docx",
        [
            "{% for item in tasks %}",
            "Opgave: {{ item.title }} - Timer: {{ item.hours }} (Nr. {{ loop.index }})",
            "{% endfor %}",
        ],
    )
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {
                        "id": "tasks",
                        "type": "repeater",
                        "label": "Opgaver",
                        "fields": [
                            {"id": "title", "type": "text", "label": "Titel"},
                            {"id": "hours", "type": "number", "label": "Timer"},
                        ],
                    }
                ],
            }
        ],
    }
    validate_docx_binding(docx_file, form_def)


def test_valid_multiselect_loop_binding(tmp_path: Path) -> None:
    docx_file = _make_docx(
        tmp_path / "document.docx",
        [
            "{% for sport in sports %}",
            "Sport: {{ sport }}",
            "{% endfor %}",
        ],
    )
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {
                        "id": "sports",
                        "type": "multiselect",
                        "label": "Sportsgrene",
                        "options": ["f", "g"],
                    }
                ],
            }
        ],
    }
    validate_docx_binding(docx_file, form_def)


def test_multiselect_loop_item_attribute_access_rejected(tmp_path: Path) -> None:
    docx_file = _make_docx(
        tmp_path / "document.docx",
        [
            "{% for sport in sports %}",
            "Sport: {{ sport.invalid_attr }}",
            "{% endfor %}",
        ],
    )
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {
                        "id": "sports",
                        "type": "multiselect",
                        "label": "Sportsgrene",
                        "options": ["f", "g"],
                    }
                ],
            }
        ],
    }
    with pytest.raises(DOCXBindingValidationError, match="Attribute access 'invalid_attr' is not allowed on multiselect loop item"):
        validate_docx_binding(docx_file, form_def)


def test_yaml_only_fields_do_not_need_to_appear_in_docx(tmp_path: Path) -> None:
    docx_file = _make_docx(tmp_path / "document.docx", ["Kun rolle: {{ role_name }}"])
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {"id": "role_name", "type": "text", "label": "Rollenavn"},
                    {"id": "privacy_info", "type": "info", "text": "Privatlivsinformation"},
                    {"id": "consent", "type": "checkbox", "label": "Samtykke", "purpose": "consent"},
                    {"id": "notes", "type": "textarea", "label": "Noter"},
                ],
            }
        ],
    }
    # Missing optional/privacy/consent fields in DOCX is valid
    validate_docx_binding(docx_file, form_def)


def test_info_field_referenced_in_docx_rejected(tmp_path: Path) -> None:
    """Top-level info fields are never document data and must not be referenced in DOCX."""
    docx_file = _make_docx(tmp_path / "document.docx", ["Info: {{ privacy_info }}"])
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {"id": "privacy_info", "type": "info", "text": "Privatliv"},
                    {"id": "role_name", "type": "text", "label": "Rollenavn"},
                ],
            }
        ],
    }
    with pytest.raises(DOCXBindingValidationError, match="references unknown variable 'privacy_info'"):
        validate_docx_binding(docx_file, form_def)


def test_unknown_root_docx_variable_rejected(tmp_path: Path) -> None:
    docx_file = _make_docx(tmp_path / "document.docx", ["Ukendt: {{ unknown_field }}"])
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [{"id": "role_name", "type": "text", "label": "Rollenavn"}],
            }
        ],
    }
    with pytest.raises(DOCXBindingValidationError, match="references unknown variable 'unknown_field'"):
        validate_docx_binding(docx_file, form_def)


def test_loop_metadata_outside_loop_rejected(tmp_path: Path) -> None:
    docx_file = _make_docx(tmp_path / "document.docx", ["Udenfor loop: {{ loop.index }}"])
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [{"id": "role_name", "type": "text", "label": "Rollenavn"}],
            }
        ],
    }
    with pytest.raises(DOCXBindingValidationError, match="Variable 'loop.index' is only valid inside a for-loop body"):
        validate_docx_binding(docx_file, form_def)


def test_standalone_loop_or_super_rejected(tmp_path: Path) -> None:
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [{"id": "role_name", "type": "text", "label": "Rollenavn"}],
            }
        ],
    }

    f1 = _make_docx(tmp_path / "super.docx", ["{{ super }}"])
    with pytest.raises(DOCXBindingValidationError, match="Variable 'super' is not permitted"):
        validate_docx_binding(f1, form_def)

    f2 = _make_docx(tmp_path / "loop_standalone.docx", ["{{ loop }}"])
    with pytest.raises(DOCXBindingValidationError, match="Invalid standalone variable 'loop'"):
        validate_docx_binding(f2, form_def)


def test_lexical_scoping_two_sequential_loops_reusing_target(tmp_path: Path) -> None:
    docx_file = _make_docx(
        tmp_path / "document.docx",
        [
            "{% for item in tasks %}",
            "Opgave: {{ item.title }}",
            "{% endfor %}",
            "{% for item in contacts %}",
            "Kontakt: {{ item.email }}",
            "{% endfor %}",
        ],
    )
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {
                        "id": "tasks",
                        "type": "repeater",
                        "label": "Opgaver",
                        "fields": [{"id": "title", "type": "text", "label": "Titel"}],
                    },
                    {
                        "id": "contacts",
                        "type": "repeater",
                        "label": "Kontakter",
                        "fields": [{"id": "email", "type": "text", "label": "E-mail"}],
                    },
                ],
            }
        ],
    }
    # Reusing 'item' across sequential loops must be valid
    validate_docx_binding(docx_file, form_def)


def test_lexical_scoping_wrong_child_in_reused_target_fails(tmp_path: Path) -> None:
    docx_file = _make_docx(
        tmp_path / "document.docx",
        [
            "{% for item in tasks %}",
            "Forkert: {{ item.email }}",
            "{% endfor %}",
            "{% for item in contacts %}",
            "Rigtig: {{ item.email }}",
            "{% endfor %}",
        ],
    )
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {
                        "id": "tasks",
                        "type": "repeater",
                        "label": "Opgaver",
                        "fields": [{"id": "title", "type": "text", "label": "Titel"}],
                    },
                    {
                        "id": "contacts",
                        "type": "repeater",
                        "label": "Kontakter",
                        "fields": [{"id": "email", "type": "text", "label": "E-mail"}],
                    },
                ],
            }
        ],
    }
    # The first loop's item is bound to tasks, which does NOT have 'email'
    with pytest.raises(DOCXBindingValidationError, match="references unknown child field 'email' on repeater 'tasks'"):
        validate_docx_binding(docx_file, form_def)


def test_output_grammar_disallows_complex_expressions(tmp_path: Path) -> None:
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {"id": "role_name", "type": "text", "label": "Rollenavn"},
                    {"id": "other_field", "type": "text", "label": "Andet"},
                    {"id": "flag", "type": "checkbox", "label": "Flag"},
                ],
            }
        ],
    }

    # 1. List construction in output
    f1 = _make_docx(tmp_path / "out_list.docx", ["{{ [role_name] }}"])
    with pytest.raises(DOCXBindingValidationError, match="Unsupported expression construct List in output tag"):
        validate_docx_binding(f1, form_def)

    # 2. Boolean 'and' in output
    f2 = _make_docx(tmp_path / "out_and.docx", ["{{ role_name and other_field }}"])
    with pytest.raises(DOCXBindingValidationError, match="Unsupported expression construct And in output tag"):
        validate_docx_binding(f2, form_def)

    # 3. Ternary conditional in output
    f3 = _make_docx(tmp_path / "out_ternary.docx", ["{{ role_name if flag else other_field }}"])
    with pytest.raises(DOCXBindingValidationError, match="Unsupported expression construct CondExpr in output tag"):
        validate_docx_binding(f3, form_def)


def test_condition_grammar_disallows_unsupported_operators(tmp_path: Path) -> None:
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {"id": "count", "type": "number", "label": "Antal"},
                ],
            }
        ],
    }

    # Greater than operator '>' is not in supported contract
    f1 = _make_docx(tmp_path / "cond_gt.docx", ["{% if count > 5 %}Mange{% endif %}"])
    with pytest.raises(DOCXBindingValidationError, match="Unsupported comparison operator 'gt'"):
        validate_docx_binding(f1, form_def)


def test_chained_attribute_access_rejected(tmp_path: Path) -> None:
    docx_file = _make_docx(
        tmp_path / "document.docx",
        [
            "{% for item in tasks %}",
            "{{ item.title.__class__ }}",
            "{% endfor %}",
        ],
    )
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {
                        "id": "tasks",
                        "type": "repeater",
                        "label": "Opgaver",
                        "fields": [{"id": "title", "type": "text", "label": "Titel"}],
                    }
                ],
            }
        ],
    }
    with pytest.raises(DOCXBindingValidationError, match="Chained or dynamic attribute access is not allowed"):
        validate_docx_binding(docx_file, form_def)


def test_subscript_getitem_access_rejected(tmp_path: Path) -> None:
    docx_file = _make_docx(
        tmp_path / "document.docx",
        [
            "{% for item in tasks %}",
            "{{ item['title'] }}",
            "{% endfor %}",
        ],
    )
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {
                        "id": "tasks",
                        "type": "repeater",
                        "label": "Opgaver",
                        "fields": [{"id": "title", "type": "text", "label": "Titel"}],
                    }
                ],
            }
        ],
    }
    with pytest.raises(DOCXBindingValidationError, match="Unsupported expression construct Getitem"):
        validate_docx_binding(docx_file, form_def)


def test_non_direct_loop_iterables_rejected(tmp_path: Path) -> None:
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {"id": "role_name", "type": "text", "label": "Rollenavn"},
                    {
                        "id": "tasks",
                        "type": "repeater",
                        "label": "Opgaver",
                        "fields": [{"id": "title", "type": "text", "label": "Titel"}],
                    },
                    {
                        "id": "other_tasks",
                        "type": "repeater",
                        "label": "Andre opgaver",
                        "fields": [{"id": "title", "type": "text", "label": "Titel"}],
                    },
                ],
            }
        ],
    }

    # 1. Filtered loop
    f1 = _make_docx(tmp_path / "f1.docx", ["{% for x in role_name | list %}{{ x }}{% endfor %}"])
    with pytest.raises(DOCXBindingValidationError, match="Loop iterable must be a direct field reference"):
        validate_docx_binding(f1, form_def)

    # 2. Arithmetic / concatenated loop
    f2 = _make_docx(tmp_path / "f2.docx", ["{% for x in tasks + other_tasks %}{{ x.title }}{% endfor %}"])
    with pytest.raises(DOCXBindingValidationError, match="Loop iterable must be a direct field reference"):
        validate_docx_binding(f2, form_def)

    # 3. If-filtered loop
    f3 = _make_docx(tmp_path / "f3.docx", ["{% for x in tasks if x.title %}{{ x.title }}{% endfor %}"])
    with pytest.raises(DOCXBindingValidationError, match="Loop filtering .* is not allowed"):
        validate_docx_binding(f3, form_def)


def test_invalid_jinja_syntax_rejected(tmp_path: Path) -> None:
    docx_file = _make_docx(tmp_path / "document.docx", ["Ugyldig: {% if role_name %} Mangler endif"])
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [{"id": "role_name", "type": "text", "label": "Rollenavn"}],
            }
        ],
    }
    with pytest.raises(DOCXBindingValidationError, match="Invalid Jinja/docxtpl syntax"):
        validate_docx_binding(docx_file, form_def)


def test_disallowed_macro_construct_rejected(tmp_path: Path) -> None:
    docx_file = _make_docx(tmp_path / "document.docx", ["{% macro my_macro() %}tekst{% endmacro %}"])
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [{"id": "role_name", "type": "text", "label": "Rollenavn"}],
            }
        ],
    }
    with pytest.raises(DOCXBindingValidationError, match="Unsupported template construct Macro"):
        validate_docx_binding(docx_file, form_def)


def test_disallowed_function_call_rejected(tmp_path: Path) -> None:
    docx_file = _make_docx(tmp_path / "document.docx", ["{{ get_data() }}"])
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [{"id": "role_name", "type": "text", "label": "Rollenavn"}],
            }
        ],
    }
    with pytest.raises(DOCXBindingValidationError, match="Unsupported expression construct Call"):
        validate_docx_binding(docx_file, form_def)


def test_corrupt_docx_file_rejected(tmp_path: Path) -> None:
    corrupt_file = tmp_path / "document.docx"
    corrupt_file.write_bytes(b"not a valid zip or docx content")
    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [{"id": "role_name", "type": "text", "label": "Rollenavn"}],
            }
        ],
    }
    with pytest.raises(DOCXBindingValidationError, match="Invalid or corrupt DOCX document template"):
        validate_docx_binding(corrupt_file, form_def)


def test_load_template_package_with_invalid_docx_binding_fails(tmp_path: Path) -> None:
    pkg_dir = tmp_path / "hif-01-role"
    pkg_dir.mkdir()
    (pkg_dir / "template.yaml").write_text(
        """
id: hif-01-role
version: 1
enabled: true
title: Rollebeskrivelse
category: Organisation
description: Beskrivelse
steps:
  - id: basic
    title: Grundoplysninger
    fields:
      - id: role_name
        type: text
        label: Rollenavn
""",
        encoding="utf-8",
    )
    _make_docx(pkg_dir / "document.docx", ["{{ unknown_injected_var }}"])

    with pytest.raises(DOCXBindingValidationError, match="references unknown variable 'unknown_injected_var'"):
        load_template_package(pkg_dir)


def test_docxtpl_structural_paragraph_and_table_row_tags(tmp_path: Path) -> None:
    """Test actual Word docxtpl structural paragraph ({%p}) and table row ({%tr}) tags."""
    doc_path = tmp_path / "structural_template.docx"
    doc = docx.Document()

    # Paragraph structural tags
    doc.add_paragraph("{%p if is_active %}")
    doc.add_paragraph("Aktiv rolle: {{ role_name }}")
    doc.add_paragraph("{%p endif %}")

    # Table structural tags
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].paragraphs[0].text = "{%tr for item in tasks %}"
    table.rows[1].cells[0].paragraphs[0].text = "{{ item.title }}"
    table.rows[1].cells[1].paragraphs[0].text = "{{ item.hours }}"
    table.rows[2].cells[0].paragraphs[0].text = "{%tr endfor %}"

    doc.save(str(doc_path))

    form_def = {
        "id": "hif-01",
        "version": 1,
        "steps": [
            {
                "id": "s1",
                "title": "Trin 1",
                "fields": [
                    {"id": "is_active", "type": "checkbox", "label": "Aktiv"},
                    {"id": "role_name", "type": "text", "label": "Rollenavn"},
                    {
                        "id": "tasks",
                        "type": "repeater",
                        "label": "Opgaver",
                        "fields": [
                            {"id": "title", "type": "text", "label": "Titel"},
                            {"id": "hours", "type": "number", "label": "Timer"},
                        ],
                    },
                ],
            }
        ],
    }

    # 1. Must pass binding validation cleanly
    validate_docx_binding(doc_path, form_def)

    # 2. Must render cleanly with docxtpl
    out_path = tmp_path / "structural_rendered.docx"
    render_docx(
        doc_path,
        {
            "is_active": True,
            "role_name": "Kasserer",
            "tasks": [{"title": "Regnskab", "hours": 12}],
        },
        out_path,
    )
    assert out_path.is_file()

    # Verify rendered output
    rendered_doc = docx.Document(str(out_path))
    body_text = " ".join(p.text for p in rendered_doc.paragraphs)
    assert "Aktiv rolle: Kasserer" in body_text

    table_cells = [c.text.strip() for row in rendered_doc.tables[0].rows for c in row.cells]
    assert "Regnskab" in table_cells
    assert "12" in table_cells
